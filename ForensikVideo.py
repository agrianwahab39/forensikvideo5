
# START OF FILE ForensikVideo.py

# vifa_pro.py
# (Sistem Forensik Video Profesional dengan Analisis Multi-Lapis)
# VERSI 5 TAHAP PENELITIAN (DENGAN PERBAIKAN BUG STYLE REPORTLAB)
# VERSI PENINGKATAN METODE UTAMA (K-MEANS, LOCALIZATION) & PENDUKUNG (ELA, SIFT)
# VERSI REVISI DETAIL TAHAP 1 (METADATA, NORMALISASI FRAME, DETAIL K-MEANS)
# VERSI PENINGKATAN DETAIL TAHAP 2 (PLOT TEMPORAL K-MEANS, SSIM, OPTICAL FLOW)
# VERSI PENINGKATAN DETAIL TAHAP 3 (INVESTIGASI MENDALAM DAN PENJELASAN LENGKAP)
# VERSI PENINGKATAN DETAIL TAHAP 4 (LOCALIZATION TAMPERING ENHANCED, SKOR INTEGRITAS REALISTIS)
# MODIFIKASI: Menghapus konsep Skor Integritas dan menggantinya sepenuhnya dengan Forensic Evidence Reliability Matrix (FERM).

"""
VIFA-Pro: Sistem Forensik Video Profesional (Arsitektur 5 Tahap)
========================================================================================
Versi ini mengimplementasikan alur kerja forensik formal dalam 5 tahap yang jelas,
sesuai dengan metodologi penelitian untuk deteksi manipulasi video. Setiap tahap
memiliki tujuan spesifik, dari ekstraksi fitur dasar hingga validasi proses.

ARSITEKTUR PIPELINE:
- TAHAP 1: Pra-pemrosesan & Ekstraksi Fitur Dasar (Hashing, Frame, pHash, Warna)
           -> Metadata diekstrak secara mendalam.
           -> Frame diekstrak dan dinormalisasi warnanya untuk konsistensi analisis.
           -> Metode K-Means diterapkan untuk klasterisasi warna adegan dengan visualisasi detail.
- TAHAP 2: Analisis Anomali Temporal & Komparatif (Optical Flow, SSIM, K-Means Temporal, Baseline Check)
           -> Visualisasi Temporal yang lebih rinci untuk SSIM, Optical Flow, dan K-Means.
- TAHAP 3: Sintesis Bukti & Investigasi Mendalam (Korelasi Metrik, ELA & SIFT on-demand)
           -> ELA dan SIFT+RANSAC digunakan sebagai investigasi pendukung yang terukur.
           -> Analisis detail dengan penjelasan lengkap untuk setiap anomali.
- TAHAP 4: Visualisasi & Penilaian Integritas (Plotting, Integrity Score)
           -> Localization Tampering menyatukan anomali menjadi peristiwa yang dapat diinterpretasikan.
           -> ENHANCED: Penilaian keandalan bukti menggunakan Forensic Evidence Reliability Matrix (FERM).
- TAHAP 5: Penyusunan Laporan & Validasi Forensik (Laporan PDF Naratif)

Deteksi:
- Diskontinuitas (Deletion/Insertion): Melalui Aliran Optik, SSIM, K-Means, dan Perbandingan Baseline.
- Duplikasi Frame (Duplication): Melalui pHash, dikonfirmasi oleh SIFT+RANSAC dan SSIM.
- Penyisipan Area (Splicing): Terindikasi oleh Analisis Tingkat Kesalahan (ELA) pada titik diskontinuitas.

Author: OpenAI-GPT & Anda
License: MIT
Dependencies: opencv-python, opencv-contrib-python, imagehash, numpy, Pillow,
              reportlab, matplotlib, tqdm, scikit-learn, scikit-image
"""

from __future__ import annotations
import argparse
import json
import hashlib
import shutil
import subprocess
import os
import zipfile
import sys
import io
import bleach
import markdown2
from jinja2 import Environment, BaseLoader
from weasyprint import HTML, CSS
from pathlib import Path
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from collections import defaultdict, Counter
from typing import Dict, List, Tuple, Optional, Any

# Pemeriksaan Dependensi Awal
try:
    import cv2
    import imagehash
    import numpy as np
    from PIL import Image, ImageChops, ImageEnhance, ImageDraw, ImageFont
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.utils import ImageReader
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as PlatypusImage, Table, TableStyle, PageBreak
    except ImportError as e:
        PlatypusImage = None
    from reportlab.lib import colors
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from tqdm import tqdm
    from sklearn.cluster import KMeans
    from skimage.metrics import structural_similarity as ssim
    from scipy import stats
    import seaborn as sns

    # Import untuk ekspor DOCX
    try:
        from docx import Document
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False
        print("Warning: python-docx tidak terinstall. Fitur ekspor DOCX tidak tersedia.")

except ImportError as e:
    print(f"Error: Dependensi penting tidak ditemukan -> {e}")
    sys.exit(1)


###############################################################################
# Utilitas & Konfigurasi Global
###############################################################################

# Konfigurasi default
CONFIG = {
    "KMEANS_CLUSTERS": 3,
    "KMEANS_SAMPLES_PER_CLUSTER": 3,
    "SSIM_DISCONTINUITY_DROP": 0.30,
    "OPTICAL_FLOW_Z_THRESH": 5.0,
    "DUPLICATION_SSIM_CONFIRM": 0.80,
    "SIFT_MIN_MATCH_COUNT": 10,
    "USE_AUTO_THRESHOLDS": True
}

class Icons:
    IDENTIFICATION = "🔍"
    PRESERVATION = "🛡️"
    COLLECTION = "📥"
    EXAMINATION = "🔬"
    ANALYSIS = "📈"
    REPORTING = "📄"
    SUCCESS = "✅"
    ERROR = "❌"
    INFO = "ℹ️"
    CONFIDENCE_LOW = "🟩"
    CONFIDENCE_MED = "🟨"
    CONFIDENCE_HIGH = "🟧"
    CONFIDENCE_VHIGH = "🟥"

# Fungsi log yang dienkapsulasi untuk output ke konsol dan UI Streamlit
def log(message: str):
    print(message, file=sys.stdout) # Menggunakan stdout asli untuk logging

def print_stage_banner(stage_number: int, stage_name: str, icon: str, description: str):
    width=80
    log("\n" + "="*width)
    log(f"=== {icon}  TAHAP {stage_number}: {stage_name.upper()} ".ljust(width - 3) + "===")
    log("="*width)
    log(f"{Icons.INFO}  {description}")
    log("-" * width)

###############################################################################
# Struktur Data Inti (DIPERLUAS UNTUK TAHAP 4 ENHANCED)
###############################################################################

@dataclass
class Evidence:
    reasons: list[str] = field(default_factory=list)
    metrics: dict = field(default_factory=dict)
    confidence: str = "N/A"
    ela_path: str | None = None
    sift_path: str | None = None
    # Tambahan untuk detail Tahap 3
    detailed_analysis: dict = field(default_factory=dict)
    visualizations: dict = field(default_factory=dict)
    explanations: dict = field(default_factory=dict)

@dataclass
class FrameInfo:
    index: int
    timestamp: float
    img_path_original: str  # Path ke frame asli
    img_path: str           # Path ke frame yang dinormalisasi (digunakan untuk analisis utama)
    img_path_comparison: str | None = None # Path ke gambar perbandingan (opsional)
    hash: str | None = None
    type: str = "original"
    ssim_to_prev: float | None = None
    optical_flow_mag: float | None = None
    color_cluster: int | None = None
    evidence_obj: Evidence = field(default_factory=Evidence)
    # Tambahan untuk analisis detail
    histogram_data: np.ndarray | None = None
    edge_density: float | None = None
    blur_metric: float | None = None

@dataclass
class AnalysisResult:
    video_path: str
    preservation_hash: str
    metadata: dict
    frames: list[FrameInfo]
    metadata_raw: dict = field(default_factory=dict)
    metadata_analysis: dict = field(default_factory=dict)
    summary: dict = field(default_factory=dict)
    plots: dict = field(default_factory=dict)
    # --- PENAMBAHAN: Artefak K-Means yang Detail ---
    kmeans_artifacts: dict = field(default_factory=dict)
    localizations: list[dict] = field(default_factory=list)
    zip_report_path: Optional[Path] = None
    pdf_report_path: Optional[Path] = None
    zip_report_bytes: Optional[bytes] = None
    pdf_report_bytes: Optional[bytes] = None

    # Tambahan untuk Tahap 3
    detailed_anomaly_analysis: dict = field(default_factory=dict)
    statistical_summary: dict = field(default_factory=dict)
    # TAMBAHAN UNTUK TAHAP 4 ENHANCED
    integrity_analysis: dict = field(default_factory=dict)  # Pastikan ini ada
    pipeline_assessment: dict = field(default_factory=dict)
    localization_details: dict = field(default_factory=dict)
    confidence_distribution: dict = field(default_factory=dict)
    # TAMBAHAN UNTUK FERM
    forensic_evidence_matrix: dict = field(default_factory=dict)

###############################################################################
# Fungsi Analisis Individual (EXISTING)
###############################################################################

def perform_ela(image_path: Path, quality: int=90) -> tuple[Path, int, np.ndarray] | None:
    """
    Error Level Analysis (ELA) yang ditingkatkan dengan analisis grid dan statistik detail.
    Mengembalikan path gambar ELA, max difference, dan array ELA untuk analisis lebih lanjut.
    """
    try:
        ela_dir = image_path.parent.parent / "ela_artifacts"
        ela_dir.mkdir(exist_ok=True)
        out_path = ela_dir / f"{image_path.stem}_ela.jpg"
        temp_jpg_path = out_path.with_name(f"temp_{out_path.name}")

        # Buka dan simpan dengan kualitas tertentu
        with Image.open(image_path).convert('RGB') as im:
            im.save(temp_jpg_path, 'JPEG', quality=quality)

        # Hitung perbedaan
        with Image.open(image_path).convert('RGB') as im_orig, Image.open(temp_jpg_path) as resaved_im:
            ela_im = ImageChops.difference(im_orig, resaved_im)

        if Path(temp_jpg_path).exists():
            Path(temp_jpg_path).unlink()

        # Konversi ke array untuk analisis statistik
        ela_array = np.array(ela_im)

        # Hitung statistik detail
        extrema = ela_im.getextrema()
        max_diff = max(ex[1] for ex in extrema) if extrema else 1
        scale = 255.0 / (max_diff if max_diff > 0 else 1)

        # Enhance dan simpan
        ela_im = ImageEnhance.Brightness(ela_im).enhance(scale)

        # Tambahkan grid untuk analisis regional
        ela_with_grid = ela_im.copy()
        draw = ImageDraw.Draw(ela_with_grid)
        width, height = ela_with_grid.size
        grid_size = 50

        for x in range(0, width, grid_size):
            draw.line([(x, 0), (x, height)], fill=(128, 128, 128), width=1)
        for y in range(0, height, grid_size):
            draw.line([(0, y), (width, y)], fill=(128, 128, 128), width=1)

        ela_with_grid.save(out_path)

        return out_path, max_diff, ela_array
    except Exception as e:
        log(f"  {Icons.ERROR} Gagal ELA pada {image_path.name}: {e}")
        return None

def analyze_ela_regions(ela_array: np.ndarray, grid_size: int = 50) -> dict:
    """
    Menganalisis ELA berdasarkan region grid untuk mendeteksi area yang mencurigakan.
    """
    height, width = ela_array.shape[:2]
    suspicious_regions = []

    for y in range(0, height, grid_size):
        for x in range(0, width, grid_size):
            # Ekstrak region
            region = ela_array[y:min(y+grid_size, height), x:min(x+grid_size, width)]
            if region.size == 0:
                continue

            # Hitung metrik untuk setiap region
            mean_val = np.mean(region)
            std_val = np.std(region)
            max_val = np.max(region)

            # Deteksi region mencurigakan (nilai ELA tinggi)
            if mean_val > 30 or max_val > 100:  # Threshold dapat disesuaikan
                suspicious_regions.append({
                    'x': x, 'y': y,
                    'width': min(grid_size, width - x),
                    'height': min(grid_size, height - y),
                    'mean_ela': float(mean_val),
                    'std_ela': float(std_val),
                    'max_ela': float(max_val),
                    'suspicion_level': 'high' if mean_val > 50 else 'medium'
                })

    return {
        'total_regions': (height // grid_size) * (width // grid_size),
        'suspicious_regions': suspicious_regions,
        'suspicious_count': len(suspicious_regions),
        'grid_size': grid_size
    }

def compare_sift_enhanced(img_path1: Path, img_path2: Path, out_dir: Path) -> dict:
    """
    SIFT comparison yang ditingkatkan dengan analisis geometri dan visualisasi detail.
    """
    try:
        img1 = cv2.imread(str(img_path1), cv2.IMREAD_GRAYSCALE)
        img2 = cv2.imread(str(img_path2), cv2.IMREAD_GRAYSCALE)
        if img1 is None or img2 is None:
            return {'success': False, 'error': 'Failed to load images'}

        # Create SIFT detector
        sift = cv2.SIFT_create()
        kp1, des1 = sift.detectAndCompute(img1, None)
        kp2, des2 = sift.detectAndCompute(img2, None)

        if des1 is None or des2 is None or len(kp1) < 2 or len(kp2) < 2:
            return {'success': False, 'error': 'Insufficient keypoints'}

        # Match features
        bf = cv2.BFMatcher()
        matches = bf.knnMatch(des1, des2, k=2)

        if not matches or any(len(m) < 2 for m in matches):
            return {'success': False, 'error': 'No valid matches'}

        # Apply ratio test
        good_matches = []
        for m, n in matches:
            if m.distance < 0.75 * n.distance:
                good_matches.append(m)

        result = {
            'success': True,
            'total_keypoints_img1': len(kp1),
            'total_keypoints_img2': len(kp2),
            'total_matches': len(matches),
            'good_matches': len(good_matches),
            'match_quality': 'excellent' if len(good_matches) > 100 else 'good' if len(good_matches) > 50 else 'fair' if len(good_matches) > 20 else 'poor'
        }

        if len(good_matches) > CONFIG["SIFT_MIN_MATCH_COUNT"]:
            # Extract matched points
            src_pts = np.float32([kp1[m.queryIdx].pt for m in good_matches]).reshape(-1, 1, 2)
            dst_pts = np.float32([kp2[m.trainIdx].pt for m in good_matches]).reshape(-1, 1, 2)

            # Find homography
            M, mask = cv2.findHomography(src_pts, dst_pts, cv2.RANSAC, 5.0)

            if M is not None and mask is not None:
                inliers = mask.ravel().sum()
                inlier_ratio = inliers / len(good_matches) if len(good_matches) > 0 else 0.0

                # Analyze transformation matrix
                det = np.linalg.det(M[:2, :2])
                scale = np.sqrt(abs(det))

                result.update({
                    'inliers': int(inliers),
                    'outliers': len(good_matches) - int(inliers),
                    'inlier_ratio': float(inlier_ratio),
                    'homography_determinant': float(det),
                    'estimated_scale': float(scale),
                    'transformation_type': 'rigid' if abs(scale - 1.0) < 0.1 else 'scaled' if 0.5 < scale < 2.0 else 'complex'
                })

                # Create detailed visualization
                draw_params = dict(
                    matchColor=(0, 255, 0),
                    singlePointColor=(255, 0, 0),
                    matchesMask=mask.ravel().tolist(),
                    flags=cv2.DrawMatchesFlags_DEFAULT
                )

                img_matches = cv2.drawMatches(img1, kp1, img2, kp2, good_matches, None, **draw_params)

                # Add text annotations
                font = cv2.FONT_HERSHEY_SIMPLEX
                cv2.putText(img_matches, f'Total Matches: {len(good_matches)}', (10, 30), font, 0.8, (255, 255, 255), 2)
                cv2.putText(img_matches, f'Inliers: {inliers} ({inlier_ratio:.1%})', (10, 60), font, 0.8, (0, 255, 0), 2)
                cv2.putText(img_matches, f'Quality: {result["match_quality"].upper()}', (10, 90), font, 0.8, (255, 255, 0), 2)

                # Save visualization
                sift_dir = out_dir / "sift_artifacts"
                sift_dir.mkdir(exist_ok=True)
                out_path = sift_dir / f"sift_detailed_{img_path1.stem}_vs_{img_path2.stem}.jpg"
                cv2.imwrite(str(out_path), img_matches)

                result['visualization_path'] = str(out_path)

                # Create heatmap of matched points
                heatmap = create_match_heatmap(src_pts, dst_pts, img1.shape, img2.shape)
                heatmap_path = sift_dir / f"sift_heatmap_{img_path1.stem}_vs_{img_path2.stem}.jpg"
                cv2.imwrite(str(heatmap_path), heatmap)
                result['heatmap_path'] = str(heatmap_path)

        return result
    except Exception as e:
        log(f"  {Icons.ERROR} Gagal SIFT: {e}")
        return {'success': False, 'error': str(e)}

def create_match_heatmap(src_pts: np.ndarray, dst_pts: np.ndarray, shape1: tuple, shape2: tuple) -> np.ndarray:
    """
    Membuat heatmap dari distribusi titik-titik yang cocok.
    """
    height = max(shape1[0], shape2[0])
    width = shape1[1] + shape2[1] + 50

    heatmap = np.zeros((height, width, 3), dtype=np.uint8)

    # Create gaussian kernel for heatmap
    kernel_size = 21
    kernel = cv2.getGaussianKernel(kernel_size, 5)
    kernel = kernel * kernel.T
    kernel = (kernel / kernel.max() * 255).astype(np.uint8)

    # Add heat for source points
    for pt in src_pts:
        x, y = int(pt[0][0]), int(pt[0][1])
        if 0 <= x < shape1[1] and 0 <= y < shape1[0]:
            cv2.circle(heatmap, (x, y), 10, (255, 0, 0), -1)

    # Add heat for destination points
    for pt in dst_pts:
        x, y = int(pt[0][0]) + shape1[1] + 50, int(pt[0][1])
        if 0 <= x < width and 0 <= y < shape2[0]:
            cv2.circle(heatmap, (x, y), 10, (0, 0, 255), -1)

    # Apply gaussian blur for smooth heatmap
    heatmap = cv2.GaussianBlur(heatmap, (31, 31), 0)

    # Apply colormap
    heatmap_gray = cv2.cvtColor(heatmap, cv2.COLOR_BGR2GRAY)
    heatmap_colored = cv2.applyColorMap(heatmap_gray, cv2.COLORMAP_JET)

    return heatmap_colored

def calculate_frame_metrics(frame_path: str) -> dict:
    """
    Menghitung metrik tambahan untuk frame: edge density, blur metric, color distribution.
    """
    try:
        img = cv2.imread(frame_path)
        if img is None:
            return {}

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Edge density menggunakan Canny
        edges = cv2.Canny(gray, 50, 150)
        edge_density = np.sum(edges > 0) / edges.size

        # Blur metric menggunakan Laplacian variance
        laplacian = cv2.Laplacian(gray, cv2.CV_64F)
        blur_metric = laplacian.var()

        # Color distribution metrics
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        hist_h = cv2.calcHist([hsv], [0], None, [180], [0, 180])
        hist_s = cv2.calcHist([hsv], [1], None, [256], [0, 256])
        hist_v = cv2.calcHist([hsv], [2], None, [256], [0, 256])

        # Normalize histograms
        if hist_h.sum() > 0: hist_h = hist_h.flatten() / hist_h.sum()
        if hist_s.sum() > 0: hist_s = hist_s.flatten() / hist_s.sum()
        if hist_v.sum() > 0: hist_v = hist_v.flatten() / hist_v.sum()

        # Calculate entropy for color diversity
        h_entropy = -np.sum(hist_h[hist_h > 0] * np.log2(hist_h[hist_h > 0]))
        s_entropy = -np.sum(hist_s[hist_s > 0] * np.log2(hist_s[hist_s > 0]))
        v_entropy = -np.sum(hist_v[hist_v > 0] * np.log2(hist_v[hist_v > 0]))

        return {
            'edge_density': float(edge_density),
            'blur_metric': float(blur_metric),
            'color_entropy': {
                'hue': float(h_entropy),
                'saturation': float(s_entropy),
                'value': float(v_entropy)
            }
        }
    except Exception as e:
        log(f"  {Icons.ERROR} Error calculating frame metrics: {e}")
        return {}

def calculate_sha256(file_path: Path) -> str:
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def ffprobe_metadata(video_path: Path) -> dict:
    cmd = ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", "-show_streams", str(video_path)]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True, encoding='utf-8')
        return json.loads(result.stdout)
    except Exception as e:
        log(f"FFprobe error: {e}")
        return {}

# --- FUNGSI BARU: PARSE METADATA DETAIL ---
def parse_ffprobe_output(metadata: dict) -> dict:
    """Mengurai output JSON ffprobe menjadi format yang lebih mudah dibaca."""
    parsed = {}
    if 'format' in metadata:
        fmt = metadata['format']
        parsed['Format'] = {
            'Filename': Path(fmt.get('filename', 'N/A')).name,
            'Format Name': fmt.get('format_long_name', 'N/A'),
            'Duration': f"{float(fmt.get('duration', 0)):.3f} s",
            'Size': f"{int(fmt.get('size', 0)) / (1024*1024):.2f} MB",
            'Bit Rate': f"{int(fmt.get('bit_rate', 0)) / 1000:.0f} kb/s",
            'Creation Time': fmt.get('tags', {}).get('creation_time', 'N/A'),
            'Encoder': fmt.get('tags', {}).get('encoder', 'N/A'),
            'Software': fmt.get('tags', {}).get('software', 'N/A'),
        }

    video_streams = [s for s in metadata.get('streams', []) if s.get('codec_type') == 'video']
    if video_streams:
        stream = video_streams[0] # Ambil stream video pertama
        parsed['Video Stream'] = {
            'Codec': stream.get('codec_name', 'N/A').upper(),
            'Profile': stream.get('profile', 'N/A'),
            'Resolution': f"{stream.get('width')}x{stream.get('height')}",
            'Aspect Ratio': stream.get('display_aspect_ratio', 'N/A'),
            'Pixel Format': stream.get('pix_fmt', 'N/A'),
            'Frame Rate': f"{eval(stream.get('r_frame_rate', '0/1')):.2f} FPS",
            'Bitrate': f"{int(stream.get('bit_rate', 0)) / 1000:.0f} kb/s" if 'bit_rate' in stream else 'N/A',
            'Encoder': stream.get('tags', {}).get('encoder', 'N/A'),
        }

    return parsed

# --- FUNGSI BARU: ANALISIS METADATA UNTUK STATUS EDITING ---
def analyze_video_metadata(metadata: dict) -> dict:
    """Menganalisis metadata terurai untuk mengidentifikasi kemungkinan jejak editing."""
    analysis = {
        'edited': False,
        'editing_software': 'Unknown',
        'creation_time': None,
    }

    fmt = metadata.get('Format', {})
    video_stream = metadata.get('Video Stream', {})
    creation_time = fmt.get('Creation Time')
    if creation_time:
        analysis['creation_time'] = creation_time

    # Cek keberadaan software/encoder yang biasa muncul akibat proses editing
    software_fields = [
        fmt.get('Encoder'),
        fmt.get('Software'),
        video_stream.get('Encoder'),
    ]
    for field in software_fields:
        if field and field != 'N/A':
            analysis['editing_software'] = field
            break

    if analysis['editing_software'] != 'Unknown':
        analysis['edited'] = True

    return analysis

# --- FUNGSI DIREVISI: EKSTRAKSI FRAME DENGAN NORMALISASI WARNA ---
def extract_frames_with_normalization(video_path: Path, out_dir: Path, fps: int) -> list[tuple[str, str, str]] | None:
    """Mengekstrak frame, menormalisasi, dan membuat gambar perbandingan."""
    original_dir = out_dir / "frames_original"
    normalized_dir = out_dir / "frames_normalized"
    comparison_dir = out_dir / "frames_comparison"
    original_dir.mkdir(parents=True, exist_ok=True)
    normalized_dir.mkdir(parents=True, exist_ok=True)
    comparison_dir.mkdir(parents=True, exist_ok=True)

    try:
        cap = cv2.VideoCapture(str(video_path))
        if not cap.isOpened():
            log(f"  {Icons.ERROR} Gagal membuka file video: {video_path}")
            return None

        # ======= [BUGFIX FPS-30] =======
        video_fps_raw = cap.get(cv2.CAP_PROP_FPS)
        # Handle cases where video_fps is 0, NaN or invalid to prevent division by zero.
        if not video_fps_raw or video_fps_raw <= 0:
            log(f"  ⚠️ Peringatan: Gagal membaca FPS video dari metadata. Menggunakan FPS asumsi (30) untuk kalkulasi.")
            video_fps = 30.0
        else:
            video_fps = video_fps_raw

        frame_paths = []
        frame_count = 0
        extracted_count = 0
        
        # Implementasi Time Accumulator untuk sampling frame yang lebih robust
        # Ini mengatasi masalah ketika `fps` (diminta) > `video_fps` (asli) dan mencegah ZeroDivisionError atau frame_skip=0
        time_increment = 1.0 / float(fps)
        next_extraction_time = 0.0
        
        pbar_total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        if pbar_total <= 0:
            pbar_total = None # Let tqdm run without a total if frame count is unknown
            
        pbar = tqdm(total=pbar_total, desc="    Ekstraksi & Normalisasi", leave=False, bar_format='{l_bar}{bar}{r_bar}')

        while cap.isOpened():
            ret, frame = cap.read()
            if not ret: break

            # Hitung waktu saat ini dalam video. Gunakan `frame_count / video_fps` sebagai fallback jika `CAP_PROP_POS_MSEC` tidak andal.
            current_time_msec = cap.get(cv2.CAP_PROP_POS_MSEC)
            if current_time_msec > 0:
                current_time = current_time_msec / 1000.0
            else:
                # Fallback jika timestamp tidak tersedia/reliabel
                current_time = frame_count / video_fps
                
            should_extract = current_time >= next_extraction_time
            # ======= [END BUGFIX] ==========

            if should_extract:
                # 1. Simpan frame original
                original_path = original_dir / f"frame_{extracted_count:06d}_orig.jpg"
                cv2.imwrite(str(original_path), frame)

                # 2. Lakukan normalisasi (Histogram Equalization pada channel Y dari YCrCb)
                ycrcb_img = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
                ycrcb_img[:, :, 0] = cv2.equalizeHist(ycrcb_img[:, :, 0])
                normalized_frame = cv2.cvtColor(ycrcb_img, cv2.COLOR_YCrCb2BGR)
                normalized_path = normalized_dir / f"frame_{extracted_count:06d}_norm.jpg"
                cv2.imwrite(str(normalized_path), normalized_frame)

                # 3. Buat gambar perbandingan
                h, w, _ = frame.shape
                comparison_img = np.zeros((h, w * 2 + 10, 3), dtype=np.uint8)
                comparison_img[:, :w] = frame
                comparison_img[:, w+10:] = normalized_frame
                cv2.putText(comparison_img, 'Original', (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
                cv2.putText(comparison_img, 'Normalized', (w + 20, 30), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
                comparison_path = comparison_dir / f"frame_{extracted_count:06d}_comp.jpg"
                cv2.imwrite(str(comparison_path), comparison_img)

                frame_paths.append((str(original_path), str(normalized_path), str(comparison_path)))
                extracted_count += 1

                # ======= [BUGFIX FPS-30] =======
                # Pindahkan accumulator ke waktu berikutnya
                next_extraction_time += time_increment
                # ======= [END BUGFIX] ==========

            frame_count += 1
            pbar.update(1)

        pbar.close()
        cap.release()
        return frame_paths
    except Exception as e:
        log(f"  {Icons.ERROR} Error saat ekstraksi frame: {e}")
        return None

###############################################################################
# FUNGSI TAMBAHAN UNTUK TAHAP 4 ENHANCED
###############################################################################

def assess_pipeline_performance(result: AnalysisResult) -> dict:
    """
    Menilai performa setiap tahap dalam pipeline forensik.
    """
    assessment = {
        'tahap_1': {
            'nama': 'Pra-pemrosesan & Ekstraksi Fitur',
            'status': 'completed',
            'quality_score': 0,
            'metrics': {},
            'issues': []
        },
        'tahap_2': {
            'nama': 'Analisis Anomali Temporal',
            'status': 'completed',
            'quality_score': 0,
            'metrics': {},
            'issues': []
        },
        'tahap_3': {
            'nama': 'Sintesis Bukti & Investigasi',
            'status': 'completed',
            'quality_score': 0,
            'metrics': {},
            'issues': []
        },
        'tahap_4': {
            'nama': 'Visualisasi & Penilaian',
            'status': 'completed',
            'quality_score': 0,
            'metrics': {},
            'issues': []
        }
    }

    # Assess Tahap 1
    if result.frames:
        total_frames = len(result.frames)
        frames_with_hash = sum(1 for f in result.frames if f.hash is not None)
        frames_with_cluster = sum(1 for f in result.frames if f.color_cluster is not None)

        assessment['tahap_1']['metrics'] = {
            'total_frames_extracted': total_frames,
            'hash_coverage': f"{frames_with_hash/total_frames*100:.1f}%",
            'clustering_coverage': f"{frames_with_cluster/total_frames*100:.1f}%",
            'metadata_completeness': len(result.metadata) > 0
        }

        quality = (frames_with_hash/total_frames + frames_with_cluster/total_frames) / 2
        assessment['tahap_1']['quality_score'] = round(quality * 100)

        if frames_with_hash < total_frames:
            assessment['tahap_1']['issues'].append('Beberapa frame gagal di-hash')

    # Assess Tahap 2
    total_frames = len(result.frames) if result.frames else 0
    frames_with_ssim = sum(1 for f in result.frames if f.ssim_to_prev is not None)
    frames_with_flow = sum(1 for f in result.frames if f.optical_flow_mag is not None)

    assessment['tahap_2']['metrics'] = {
        'ssim_coverage': f"{frames_with_ssim/total_frames*100:.1f}%" if total_frames > 0 else "0%",
        'optical_flow_coverage': f"{frames_with_flow/total_frames*100:.1f}%" if total_frames > 0 else "0%",
        'temporal_metrics_computed': frames_with_ssim > 0 and frames_with_flow > 0
    }

    if total_frames > 0:
        quality = (frames_with_ssim + frames_with_flow) / (2 * total_frames) if total_frames > 0 else 0
        assessment['tahap_2']['quality_score'] = round(quality * 100)

    # Assess Tahap 3
    anomaly_count = sum(1 for f in result.frames if f.type.startswith('anomaly'))
    evidence_count = sum(1 for f in result.frames if f.evidence_obj.reasons)

    assessment['tahap_3']['metrics'] = {
        'anomalies_detected': anomaly_count,
        'evidence_collected': evidence_count,
        'ela_analyses': sum(1 for f in result.frames if f.evidence_obj.ela_path is not None),
        'sift_analyses': sum(1 for f in result.frames if f.evidence_obj.sift_path is not None)
    }

    if evidence_count > 0:
        assessment['tahap_3']['quality_score'] = min(100, round(evidence_count / anomaly_count * 100)) if anomaly_count > 0 else 100

    # Assess Tahap 4
    assessment['tahap_4']['metrics'] = {
        'localizations_created': len(result.localizations),
        'plots_generated': len(result.plots),
        'integrity_calculated': 'forensic_evidence_matrix' in result.__dict__
    }

    assessment['tahap_4']['quality_score'] = 100 if result.localizations else 0

    return assessment

def create_enhanced_localization_map(result: AnalysisResult, out_dir: Path) -> Path:
    """
    Membuat peta lokalisasi tampering yang lebih detail dengan timeline visual.
    """
    fig = plt.figure(figsize=(20, 12))

    # Create grid layout
    gs = fig.add_gridspec(4, 3, height_ratios=[1, 2, 1, 1], hspace=0.3, wspace=0.2)

    # Title and header
    ax_title = fig.add_subplot(gs[0, :])
    ax_title.text(0.5, 0.5, 'PETA DETAIL LOKALISASI TAMPERING',
                  ha='center', va='center', fontsize=20, weight='bold')
    ax_title.axis('off')

    # Main timeline plot
    ax_timeline = fig.add_subplot(gs[1, :])

    # Setup timeline
    total_frames = len(result.frames)
    frame_indices = list(range(total_frames))

    # Create background
    ax_timeline.axhspan(0, 1, facecolor='lightgreen', alpha=0.3, label='Normal')

    # Plot anomalies with different heights and colors
    anomaly_types = {
        'anomaly_duplication': {'color': '#FF6B6B', 'height': 0.8, 'label': 'Duplikasi', 'marker': 'o'},
        'anomaly_insertion': {'color': '#4ECDC4', 'height': 0.7, 'label': 'Penyisipan', 'marker': 's'},
        'anomaly_discontinuity': {'color': '#45B7D1', 'height': 0.6, 'label': 'Diskontinuitas', 'marker': '^'}
    }

    # Draw localization events
    for loc in result.localizations:
        event_type = loc['event']
        if event_type in anomaly_types:
            style = anomaly_types[event_type]
            start_idx = loc['start_frame']
            end_idx = loc['end_frame']

            # Draw rectangle for event duration
            rect = plt.Rectangle((start_idx, 0), end_idx - start_idx + 1, style['height'],
                               facecolor=style['color'], alpha=0.6, edgecolor='black', linewidth=2)
            ax_timeline.add_patch(rect)

            # Add confidence indicator
            conf_y = style['height'] + 0.05
            conf_color = 'red' if loc['confidence'] == 'SANGAT TINGGI' else 'orange' if loc['confidence'] == 'TINGGI' else 'yellow'
            ax_timeline.plot((start_idx + end_idx) / 2, conf_y, marker='*',
                           markersize=15, color=conf_color, markeredgecolor='black')

    # Timeline settings
    ax_timeline.set_xlim(0, total_frames)
    ax_timeline.set_ylim(0, 1)
    ax_timeline.set_xlabel('Indeks Frame', fontsize=14)
    ax_timeline.set_title('Timeline Anomali Terdeteksi', fontsize=16, pad=20)

    # Add legend
    legend_elements = [plt.Rectangle((0, 0), 1, 1, fc=style['color'], alpha=0.6, label=style['label'])
                      for style in anomaly_types.values()]
    legend_elements.append(plt.Line2D([0], [0], marker='*', color='red', markersize=10,
                                    label='Kepercayaan Tinggi', linestyle='None'))
    ax_timeline.legend(handles=legend_elements, loc='upper right', fontsize=12)

    # Add grid
    ax_timeline.grid(True, axis='x', alpha=0.3)

    # Statistics panel
    ax_stats = fig.add_subplot(gs[2, 0])
    stats_text = f"""STATISTIK ANOMALI

Total Frame: {total_frames}
Anomali Terdeteksi: {sum(1 for f in result.frames if f.type.startswith('anomaly'))}
Peristiwa Terlokalisasi: {len(result.localizations)}

Distribusi Kepercayaan:
- Sangat Tinggi: {sum(1 for loc in result.localizations if loc['confidence'] == 'SANGAT TINGGI')}
- Tinggi: {sum(1 for loc in result.localizations if loc['confidence'] == 'TINGGI')}
- Sedang: {sum(1 for loc in result.localizations if loc['confidence'] == 'SEDANG')}
- Rendah: {sum(1 for loc in result.localizations if loc['confidence'] == 'RENDAH')}"""

    ax_stats.text(0.05, 0.95, stats_text, transform=ax_stats.transAxes,
                 fontsize=11, verticalalignment='top', fontfamily='monospace',
                 bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.8))
    ax_stats.axis('off')

    # Event details panel
    ax_details = fig.add_subplot(gs[2, 1:])
    details_text = "DETAIL PERISTIWA SIGNIFIKAN\n\n"

    # Find most significant events
    significant_events = sorted(result.localizations,
                              key=lambda x: (x.get('confidence') == 'SANGAT TINGGI',
                                           x['end_frame'] - x['start_frame']),
                              reverse=True)[:5]

    for i, event in enumerate(significant_events):
        event_type = event['event'].replace('anomaly_', '').capitalize()
        duration = event['end_ts'] - event['start_ts']
        details_text += f"{i+1}. {event_type} @ {event['start_ts']:.1f}s-{event['end_ts']:.1f}s "
        details_text += f"(Durasi: {duration:.1f}s, Kepercayaan: {event.get('confidence', 'N/A')})\n"

    ax_details.text(0.05, 0.95, details_text, transform=ax_details.transAxes,
                   fontsize=11, verticalalignment='top',
                   bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.8))
    ax_details.axis('off')

    # Confidence distribution pie chart
    ax_pie = fig.add_subplot(gs[3, 0])
    confidence_counts = Counter(loc.get('confidence', 'N/A') for loc in result.localizations)
    if confidence_counts:
        colors_conf = {'SANGAT TINGGI': '#FF0000', 'TINGGI': '#FFA500',
                      'SEDANG': '#FFFF00', 'RENDAH': '#00FF00', 'N/A': '#808080'}
        pie_colors = [colors_conf.get(conf, '#808080') for conf in confidence_counts.keys()]
        ax_pie.pie(confidence_counts.values(), labels=list(confidence_counts.keys()),
                  colors=pie_colors, autopct='%1.1f%%', startangle=90)
        ax_pie.set_title('Distribusi Tingkat Kepercayaan', fontsize=12)
    else:
        ax_pie.text(0.5, 0.5, 'Tidak ada anomali', ha='center', va='center')
        ax_pie.set_xlim(0, 1)
        ax_pie.set_ylim(0, 1)

    # Temporal clustering visualization
    ax_cluster = fig.add_subplot(gs[3, 1:])

    # Calculate temporal density
    window_size = total_frames // 20 if total_frames > 20 else 1
    density = np.zeros(total_frames)

    for f in result.frames:
        if f.type.startswith('anomaly'):
            start = max(0, f.index - window_size // 2)
            end = min(total_frames, f.index + window_size // 2)
            density[start:end] += 1

    ax_cluster.fill_between(frame_indices, density, alpha=0.5, color='red')
    ax_cluster.set_xlabel('Indeks Frame', fontsize=12)
    ax_cluster.set_ylabel('Kepadatan Anomali', fontsize=12)
    ax_cluster.set_title('Analisis Kepadatan Temporal Anomali', fontsize=12)
    ax_cluster.grid(True, alpha=0.3)

    # Save the enhanced map
    enhanced_map_path = out_dir / f"enhanced_localization_map_{Path(result.video_path).stem}.png"
    plt.savefig(enhanced_map_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return enhanced_map_path

def create_anomaly_explanation_infographic(result: AnalysisResult, out_dir: Path) -> Path:
    """
    Membuat infografis yang menjelaskan setiap jenis anomali untuk orang awam.
    """
    fig = plt.figure(figsize=(16, 10))
    fig.suptitle('PANDUAN MEMAHAMI ANOMALI VIDEO', fontsize=20, fontweight='bold')

    # Define anomaly types with explanations
    anomaly_info = {
        'Duplikasi': {
            'icon': '🔁',
            'color': '#FF6B6B',
            'simple': 'Frame yang sama diulang beberapa kali',
            'technical': 'Deteksi melalui perbandingan hash dan SIFT',
            'implication': 'Bisa untuk memperpanjang durasi atau menyembunyikan penghapusan',
            'example': 'Seperti memfotokopi halaman yang sama beberapa kali'
        },
        'Diskontinuitas': {
            'icon': '✂️',
            'color': '#45B7D1',
            'simple': 'Terjadi "lompatan" atau patahan dalam aliran video',
            'technical': 'Terdeteksi melalui penurunan SSIM dan lonjakan optical flow',
            'implication': 'Indikasi pemotongan atau penyambungan yang kasar',
            'example': 'Seperti halaman yang hilang dalam sebuah buku'
        },
        'Penyisipan': {
            'icon': '➕',
            'color': '#4ECDC4',
            'simple': 'Frame baru yang tidak ada di video asli',
            'technical': 'Terdeteksi melalui perbandingan dengan baseline',
            'implication': 'Konten tambahan yang mungkin mengubah narasi',
            'example': 'Seperti menambahkan halaman baru ke dalam buku'
        }
    }

    # Create grid for each anomaly type
    gs = fig.add_gridspec(len(anomaly_info), 1, hspace=0.3, wspace=0.2)

    for idx, (atype, info) in enumerate(anomaly_info.items()):
        ax = fig.add_subplot(gs[idx])

        # Background color
        ax.add_patch(plt.Rectangle((0, 0), 1, 1, transform=ax.transAxes,
                                  facecolor=info['color'], alpha=0.1, zorder=0))

        # Title with icon
        ax.text(0.02, 0.85, f"{info['icon']} {atype.upper()}",
               transform=ax.transAxes, fontsize=18, fontweight='bold',
               bbox=dict(boxstyle='round', facecolor=info['color'], alpha=0.3))

        # Simple explanation
        ax.text(0.02, 0.65, f"Apa itu?", transform=ax.transAxes,
               fontsize=12, fontweight='bold')
        ax.text(0.02, 0.45, info['simple'], transform=ax.transAxes,
               fontsize=11, wrap=True, va='top')

        # Example
        ax.text(0.02, 0.25, f"Analogi:", transform=ax.transAxes,
               fontsize=12, fontweight='bold')
        ax.text(0.02, 0.05, info['example'], transform=ax.transAxes,
               fontsize=11, fontstyle='italic', va='top')

        # Technical
        ax.text(0.52, 0.65, f"Cara Deteksi:", transform=ax.transAxes,
               fontsize=12, fontweight='bold')
        ax.text(0.52, 0.45, info['technical'], transform=ax.transAxes,
               fontsize=11, va='top')

        # Implication
        ax.text(0.52, 0.25, f"Implikasi:", transform=ax.transAxes,
               fontsize=12, fontweight='bold')
        ax.text(0.52, 0.05, info['implication'], transform=ax.transAxes,
               fontsize=11, va='top')

        # Count from actual data
        count = sum(1 for loc in result.localizations
                   if atype.lower() in loc.get('event', '').lower())
        ax.text(0.98, 0.85, f"Ditemukan: {count}", transform=ax.transAxes,
               fontsize=14, ha='right', fontweight='bold',
               bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))

        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')

    plt.tight_layout(rect=[0, 0, 1, 0.96])
    # Save
    infographic_path = out_dir / f"anomaly_explanation_{Path(result.video_path).stem}.png"
    plt.savefig(infographic_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return infographic_path

def generate_forensic_evidence_matrix(result: AnalysisResult) -> dict:
    """
    Menghasilkan Forensic Evidence Reliability Matrix (FERM) yang menilai bukti
    forensik dari berbagai dimensi sebagai alternatif yang lebih ilmiah dibanding
    skor integritas tunggal.
    
    Returns:
        dict: Matriks hasil analisis multi-dimensi
    """
    # Inisialisasi matriks FERM
    ferm = {
        'evidence_strength': {
            'multi_method_confirmation': {},
            'confidence_distribution': {},
            'false_positive_assessment': {}
        },
        'anomaly_characterization': {
            'temporal_distribution': {},
            'technical_severity': {},
            'semantic_context': {}
        },
        'causality_analysis': {
            'technical_causes': {},
            'compression_vs_manipulation': {},
            'alternative_explanations': {}
        },
        'conclusion': {
            'primary_findings': [],
            'reliability_assessment': '',
            'recommended_actions': []
        }
    }
    
    # 1. Analisis Kekuatan Bukti
    
    # 1.1 Multi-method confirmation (berapa metode independen mengkonfirmasi anomali yang sama)
    method_confirmations = defaultdict(set)
    
    for f in result.frames:
        if not f.type.startswith('anomaly'):
            continue
            
        if f.evidence_obj.reasons:
            reasons = f.evidence_obj.reasons.split(', ') if isinstance(f.evidence_obj.reasons, str) else f.evidence_obj.reasons
            for reason in reasons:
                if "SSIM" in reason:
                    method_confirmations[f.index].add('ssim')
                if "Aliran Optik" in reason:
                    method_confirmations[f.index].add('optical_flow')
                if "Adegan" in reason or "K-Means" in reason:
                    method_confirmations[f.index].add('kmeans')
                if "ELA" in reason:
                    method_confirmations[f.index].add('ela')
                if "SIFT" in reason or "duplikasi" in reason.lower():
                    method_confirmations[f.index].add('sift')
    
    # Hitung distribusi konfirmasi multi-metode
    method_confirmation_counts = Counter([len(methods) for methods in method_confirmations.values()])
    ferm['evidence_strength']['multi_method_confirmation'] = {
        'counts': dict(method_confirmation_counts),
        'average_methods_per_anomaly': sum(k*v for k,v in method_confirmation_counts.items()) / sum(method_confirmation_counts.values()) if method_confirmation_counts else 0,
        'max_methods': max(method_confirmation_counts.keys()) if method_confirmation_counts else 0,
        'percentage_confirmed_by_multiple': sum(method_confirmation_counts[k] for k in method_confirmation_counts if k > 1) / sum(method_confirmation_counts.values()) if method_confirmation_counts else 0
    }
    
    # 1.2 Confidence distribution
    if hasattr(result, 'confidence_distribution') and result.confidence_distribution:
        ferm['evidence_strength']['confidence_distribution'] = result.confidence_distribution
    else:
        confidence_levels = Counter()
        for f in result.frames:
            if f.type.startswith('anomaly') and hasattr(f.evidence_obj, 'confidence'):
                confidence_levels[f.evidence_obj.confidence] += 1
        ferm['evidence_strength']['confidence_distribution'] = dict(confidence_levels)
    
    # 1.3 False positive assessment
    # Perkirakan probabilitas false positive berdasarkan kekuatan bukti
    false_positive_risks = {
        'SANGAT TINGGI': 0.05,  # 5% chance of false positive
        'TINGGI': 0.15,         # 15% chance of false positive
        'SEDANG': 0.30,         # 30% chance of false positive
        'RENDAH': 0.50          # 50% chance of false positive
    }
    
    if hasattr(result, 'confidence_distribution') and result.confidence_distribution:
        weighted_fp_risk = sum(
            count * false_positive_risks.get(level, 0.5) 
            for level, count in result.confidence_distribution.items()
        ) / sum(result.confidence_distribution.values()) if sum(result.confidence_distribution.values()) > 0 else 0.5
    else:
        weighted_fp_risk = 0.25  # Default value
    
    ferm['evidence_strength']['false_positive_assessment'] = {
        'weighted_risk': weighted_fp_risk,
        'risk_factors': identify_false_positive_risk_factors(result),
        'reliability_score': 1.0 - weighted_fp_risk
    }
    
    # 2. Karakterisasi Anomali
    
    # 2.1 Temporal distribution
    anomaly_frames = [f.index for f in result.frames if f.type.startswith('anomaly')]
    total_frames = len(result.frames)
    
    # Hitung kluster temporal dengan menggunakan definisi: anomali yang berjarak < 3 frame
    temporal_clusters = []
    if anomaly_frames:
        current_cluster = [anomaly_frames[0]]
        for i in range(1, len(anomaly_frames)):
            if anomaly_frames[i] - anomaly_frames[i-1] <= 3:  # Part of current cluster
                current_cluster.append(anomaly_frames[i])
            else:  # Start new cluster
                if current_cluster:
                    temporal_clusters.append(current_cluster)
                current_cluster = [anomaly_frames[i]]
        if current_cluster:  # Add the last cluster
            temporal_clusters.append(current_cluster)
    
    ferm['anomaly_characterization']['temporal_distribution'] = {
        'total_anomalies': len(anomaly_frames),
        'anomaly_density': len(anomaly_frames) / total_frames if total_frames > 0 else 0,
        'cluster_count': len(temporal_clusters),
        'avg_cluster_size': np.mean([len(c) for c in temporal_clusters]) if temporal_clusters else 0,
        'largest_cluster': max([len(c) for c in temporal_clusters]) if temporal_clusters else 0,
        'distribution_pattern': 'terisolasi' if len(temporal_clusters) > len(anomaly_frames) * 0.7 else 
                               'terkelompok' if len(temporal_clusters) > 1 else 
                               'sistematis' if len(anomaly_frames) > 0 else 'tidak ada'
    }
    
    # 2.2 Technical severity
    # Kelompokkan anomali berdasarkan tipe dan hitung severity masing-masing
    severity_by_type = defaultdict(list)
    for f in result.frames:
        if not f.type.startswith('anomaly'):
            continue
        
        severity = 0.0
        # Calculate severity based on available metrics
        if hasattr(f.evidence_obj, 'metrics'):
            metrics = f.evidence_obj.metrics
            if isinstance(metrics, dict):
                if 'ssim_drop' in metrics:
                    severity += min(1.0, metrics['ssim_drop'] * 2)  # Scale: 0.5 drop = 1.0 severity
                if 'optical_flow_z_score' in metrics:
                    severity += min(1.0, abs(metrics['optical_flow_z_score']) / 10.0)  # Scale: z-score of 10 = 1.0 severity
                if 'sift_inlier_ratio' in metrics:
                    severity += metrics['sift_inlier_ratio']  # Already 0-1 scale
                if 'ela_max_difference' in metrics:
                    severity += min(1.0, metrics['ela_max_difference'] / 200.0)  # Scale: 200 diff = 1.0 severity
                
        # Normalize severity to 0-1 range
        if 'metrics' in f.evidence_obj.__dict__ and f.evidence_obj.metrics:
            severity = min(1.0, severity / len(f.evidence_obj.metrics)) if len(f.evidence_obj.metrics) > 0 else 0.0
        else:
            # Fallback based on confidence level
            confidence_severity = {
                'SANGAT TINGGI': 0.9,
                'TINGGI': 0.7,
                'SEDANG': 0.5,
                'RENDAH': 0.3,
                'N/A': 0.1
            }
            severity = confidence_severity.get(f.evidence_obj.confidence, 0.3)
            
        severity_by_type[f.type].append(severity)
    
    ferm['anomaly_characterization']['technical_severity'] = {
        'by_type': {k: {'mean': np.mean(v), 'max': max(v), 'count': len(v)} 
                    for k, v in severity_by_type.items() if v},
        'overall_mean_severity': np.mean([s for sublist in severity_by_type.values() for s in sublist]) 
                                if any(severity_by_type.values()) else 0.0,
        'high_severity_count': sum(1 for sublist in severity_by_type.values() 
                                 for s in sublist if s > 0.7),
        'severity_distribution': Counter([('tinggi' if s > 0.7 else 
                                         'sedang' if s > 0.4 else 'rendah') 
                                         for sublist in severity_by_type.values() 
                                         for s in sublist])
    }
    
    # 2.3 Semantic context
    # Untuk analisis semantik lengkap memerlukan integrasi dengan model CV atau metode lain
    # Implementasi sederhana berdasarkan data yang tersedia
    anomaly_events = {}
    if hasattr(result, 'localizations') and result.localizations:
        for loc in result.localizations:
            event_type = loc.get('event', '').replace('anomaly_', '')
            anomaly_events[event_type] = anomaly_events.get(event_type, 0) + 1
    
    ferm['anomaly_characterization']['semantic_context'] = {
        'event_types': anomaly_events,
        'significant_events': len([loc for loc in result.localizations 
                                  if loc.get('severity_score', 0) > 0.7]) 
                             if hasattr(result, 'localizations') else 0,
        'content_analysis': 'Memerlukan integrasi model content-aware'
    }
    
    # 3. Analisis Kausalitas
    
    # 3.1 Technical causes
    tech_causes = analyze_technical_causes(result)
    ferm['causality_analysis']['technical_causes'] = tech_causes
    
    # 3.2 Compression vs manipulation
    compression_analysis = analyze_compression_artifacts(result)
    ferm['causality_analysis']['compression_vs_manipulation'] = compression_analysis
    
    # 3.3 Alternative explanations
    ferm['causality_analysis']['alternative_explanations'] = generate_alternative_explanations(result)
    
    # 4. Conclusion
    # Generate key findings and recommendations based on the analysis
    ferm['conclusion'] = generate_forensic_conclusions(result, ferm)
    
    return ferm

def identify_false_positive_risk_factors(result: AnalysisResult) -> list:
    """
    Mengidentifikasi faktor-faktor yang dapat meningkatkan risiko false positive.
    """
    risk_factors = []
    
    # Cek kualitas video dari metadata
    if hasattr(result, 'metadata'):
        video_stream = result.metadata.get('Video Stream', {})
        
        # Check for low bitrate
        bitrate_str = video_stream.get('Bitrate', 'N/A')
        if bitrate_str != 'N/A':
            try:
                bitrate = float(bitrate_str.split()[0])
                if bitrate < 500:  # Less than 500 kbps
                    risk_factors.append({
                        'factor': 'Bitrate Rendah',
                        'value': bitrate_str,
                        'impact': 'Kompresi tinggi dapat menyebabkan artefak yang mirip dengan manipulasi'
                    })
            except (ValueError, IndexError):
                pass
        
        # Check for highly compressed formats
        codec = video_stream.get('Codec', 'N/A')
        if codec in ['MPEG-4', 'H.264'] and 'Bitrate Rendah' in [rf['factor'] for rf in risk_factors]:
            risk_factors.append({
                'factor': 'Format Kompresi Tinggi',
                'value': codec,
                'impact': 'Artefak kompresi dapat salah diidentifikasi sebagai tampering'
            })
    
    # Check for very short video
    if len(result.frames) < 30:
        risk_factors.append({
            'factor': 'Durasi Video Pendek',
            'value': f'{len(result.frames)} frame',
            'impact': 'Ukuran sampel terbatas meningkatkan ketidakpastian statistik'
        })
    
    # Check for inconsistent frame rate
    fps_issues = check_frame_rate_consistency(result)
    if fps_issues:
        risk_factors.append({
            'factor': 'Frame Rate Tidak Konsisten',
            'value': fps_issues,
            'impact': 'Dapat menyebabkan positif palsu dalam analisis temporal'
        })
    
    # Check for too many anomalies (may indicate false positives)
    anomaly_count = sum(1 for f in result.frames if f.type.startswith('anomaly'))
    if anomaly_count > len(result.frames) * 0.3:  # More than 30% of frames flagged
        risk_factors.append({
            'factor': 'Deteksi Anomali Berlebihan',
            'value': f'{anomaly_count}/{len(result.frames)} frame ({anomaly_count/len(result.frames)*100:.1f}%)',
            'impact': 'Proporsi frame yang ditandai tinggi menunjukkan kemungkinan positif palsu'
        })
    
    return risk_factors

def check_frame_rate_consistency(result: AnalysisResult) -> str:
    """
    Memeriksa konsistensi frame rate dalam video.
    """
    # This would require temporal analysis of frame timestamps
    # Simplified implementation for concept demonstration
    return "Analisis memerlukan data timestamp frame yang detail"

def analyze_technical_causes(result: AnalysisResult) -> dict:
    """
    Menganalisis kemungkinan penyebab teknis untuk anomali yang terdeteksi.
    """
    causes = {
        'duplication': {
            'cause': 'Duplikasi frame',
            'technical_indicators': ['Nilai hash identik', 'Jumlah kecocokan SIFT tinggi', 'Nilai SSIM mendekati 1.0'],
            'probability': 'Tinggi' if any(f.type == 'anomaly_duplication' for f in result.frames) else 'Rendah'
        },
        'discontinuity': {
            'cause': 'Penghapusan atau penyisipan frame',
            'technical_indicators': ['Penurunan SSIM', 'Lonjakan aliran optik', 'Perubahan adegan mendadak'],
            'probability': 'Tinggi' if any(f.type == 'anomaly_discontinuity' for f in result.frames) else 'Rendah'
        },
        'insertion': {
            'cause': 'Penyambungan konten (splicing)',
            'technical_indicators': ['Anomali ELA', 'Artefak kompresi tidak konsisten', 'Ketidakcocokan dengan baseline'],
            'probability': 'Tinggi' if any(f.type == 'anomaly_insertion' for f in result.frames) else 'Rendah'
        }
    }
    
    # Count instances of each anomaly type
    type_counts = Counter(f.type for f in result.frames if f.type.startswith('anomaly'))
    
    # Add counts to the causes dictionary
    for anomaly_type, count in type_counts.items():
        clean_type = anomaly_type.replace('anomaly_', '')
        if clean_type in causes:
            causes[clean_type]['count'] = count
            causes[clean_type]['percentage'] = count / len(result.frames) * 100 if result.frames else 0
    
    return causes

def analyze_compression_artifacts(result: AnalysisResult) -> dict:
    """
    Analisis untuk membedakan antara artefak kompresi normal dan manipulasi.
    """
    # Count frames with ELA evidence
    ela_evidence_count = sum(1 for f in result.frames 
                           if f.type.startswith('anomaly') and 
                           f.evidence_obj.ela_path is not None)
    
    # Examine ELA patterns across the video
    ela_pattern = 'konsisten' if ela_evidence_count < len(result.frames) * 0.1 else 'bervariasi'
    
    # Look at compression information from metadata
    compression_info = "Tidak diketahui"
    if hasattr(result, 'metadata') and 'Video Stream' in result.metadata:
        codec = result.metadata['Video Stream'].get('Codec', 'Unknown')
        bitrate = result.metadata['Video Stream'].get('Bitrate', 'Unknown')
        compression_info = f"{codec} pada {bitrate}"
    
    return {
        'compression_info': compression_info,
        'ela_evidence_pattern': ela_pattern,
        'ela_evidence_count': ela_evidence_count,
        'compression_vs_manipulation_assessment': 
            'Kemungkinan besar manipulasi' if ela_evidence_count > 10 and ela_pattern == 'bervariasi' else
            'Mungkin manipulasi' if ela_evidence_count > 5 else
            'Kemungkinan besar artefak kompresi normal' if ela_evidence_count <= 5 else
            'Tidak dapat disimpulkan'
    }

def generate_alternative_explanations(result: AnalysisResult) -> dict:
    """
    Menghasilkan penjelasan alternatif untuk anomali yang terdeteksi.
    """
    alternatives = {
        'compression_artifacts': {
            'explanation': 'Artefak kompresi normal dapat menyebabkan anomali ELA',
            'affected_methods': ['ELA'],
            'likelihood': 'Sedang',
            'distinguishing_factors': 'Pola ELA yang konsisten di seluruh video menunjukkan kompresi daripada manipulasi yang ditargetkan'
        },
        'scene_transitions': {
            'explanation': 'Perubahan adegan normal dapat memicu penurunan SSIM dan lonjakan aliran optik',
            'affected_methods': ['SSIM', 'Optical Flow'],
            'likelihood': 'Sedang-Tinggi',
            'distinguishing_factors': 'Perubahan adegan biasanya menunjukkan perubahan warna/konten dan perubahan gerakan secara bersamaan'
        },
        'camera_movement': {
            'explanation': 'Gerakan kamera yang cepat dapat menyebabkan anomali aliran optik',
            'affected_methods': ['Optical Flow'],
            'likelihood': 'Sedang',
            'distinguishing_factors': 'Gerakan kamera biasanya mempengaruhi seluruh frame secara konsisten'
        },
        'lighting_changes': {
            'explanation': 'Perubahan pencahayaan mendadak dapat memicu pergeseran klaster K-means',
            'affected_methods': ['K-means'],
            'likelihood': 'Sedang',
            'distinguishing_factors': 'Perubahan pencahayaan mempengaruhi kecerahan keseluruhan tanpa mengubah struktur konten'
        },
        'repeated_content': {
            'explanation': 'Konten yang berulang secara alami dapat memicu deteksi duplikasi',
            'affected_methods': ['pHash', 'SIFT'],
            'likelihood': 'Rendah',
            'distinguishing_factors': 'Pengulangan alami biasanya menunjukkan variasi kecil dalam nilai piksel yang tepat'
        }
    }
    
    # Analyze which alternative explanations are most relevant for this video
    relevant_alternatives = {}
    
    # Check for potential compression artifacts
    if hasattr(result, 'metadata') and 'Video Stream' in result.metadata:
        bitrate_str = result.metadata['Video Stream'].get('Bitrate', 'N/A')
        if bitrate_str != 'N/A':
            try:
                bitrate = float(bitrate_str.split()[0])
                if bitrate < 1000:  # Less than 1 Mbps
                    alternatives['compression_artifacts']['likelihood'] = 'Tinggi'
                    relevant_alternatives['compression_artifacts'] = alternatives['compression_artifacts']
            except (ValueError, IndexError):
                pass
    
    # Check for scene transitions
    kmeans_changes = sum(1 for i in range(1, len(result.frames)) 
                        if (result.frames[i].color_cluster is not None and 
                            result.frames[i-1].color_cluster is not None and 
                            result.frames[i].color_cluster != result.frames[i-1].color_cluster))
    
    if kmeans_changes > 0:
        alternatives['scene_transitions']['likelihood'] = 'Tinggi'
        relevant_alternatives['scene_transitions'] = alternatives['scene_transitions']
    
    # Check for potential camera movement
    high_flow_frames = sum(1 for f in result.frames 
                          if f.optical_flow_mag is not None and f.optical_flow_mag > 1.0)
    
    if high_flow_frames > len(result.frames) * 0.1:  # More than 10% of frames have high flow
        alternatives['camera_movement']['likelihood'] = 'Tinggi'
        relevant_alternatives['camera_movement'] = alternatives['camera_movement']
    
    # Add other alternatives with medium or higher likelihood
    for key, alt in alternatives.items():
        if key not in relevant_alternatives and alt['likelihood'] in ['Sedang', 'Sedang-Tinggi', 'Tinggi']:
            relevant_alternatives[key] = alt
    
    return {
        'all_alternatives': alternatives,
        'relevant_alternatives': relevant_alternatives,
        'most_likely_alternative': max(relevant_alternatives.items(), 
                                     key=lambda x: {'Rendah': 1, 'Sedang': 2, 'Sedang-Tinggi': 3, 'Tinggi': 4}[x[1]['likelihood']])[0]
                                     if relevant_alternatives else None
    }

def generate_forensic_conclusions(result: AnalysisResult, ferm: dict) -> dict:
    """
    Menghasilkan kesimpulan dan rekomendasi berdasarkan analisis FERM.
    MODIFIKASI: Semua string output diterjemahkan ke Bahasa Indonesia.
    """
    # Extract key metrics for decision making
    evidence_strength = ferm['evidence_strength']
    anomaly_char = ferm['anomaly_characterization']
    causality = ferm['causality_analysis']
    
    # Calculate overall confidence
    confidence_weights = {
        'SANGAT TINGGI': 4,
        'TINGGI': 3,
        'SEDANG': 2,
        'RENDAH': 1
    }
    
    confidence_dist = evidence_strength['confidence_distribution']
    if confidence_dist:
        weighted_confidence = sum(confidence_weights.get(level, 0) * count 
                                for level, count in confidence_dist.items())
        total_anomalies = sum(confidence_dist.values())
        avg_confidence = weighted_confidence / total_anomalies if total_anomalies > 0 else 0
    else:
        avg_confidence = 0
    
    # Determine primary findings
    primary_findings = []
    
    # Check for duplication events
    duplication_count = sum(1 for f in result.frames if f.type == 'anomaly_duplication')
    if duplication_count > 0:
        primary_findings.append({
            'finding': f"Terdeteksi {duplication_count} frame duplikat",
            'confidence': 'Tinggi' if avg_confidence > 3 else 'Sedang' if avg_confidence > 2 else 'Rendah',
            'evidence': 'Kecocokan hash, konfirmasi SIFT, skor SSIM tinggi',
            'interpretation': 'Mengindikasikan potensi manipulasi untuk memperpanjang durasi atau menyembunyikan penghapusan konten'
        })
    
    # Check for discontinuity events
    discontinuity_count = sum(1 for f in result.frames if f.type == 'anomaly_discontinuity')
    if discontinuity_count > 0:
        primary_findings.append({
            'finding': f"Terdeteksi {discontinuity_count} frame dengan diskontinuitas temporal",
            'confidence': 'Tinggi' if avg_confidence > 3 else 'Sedang' if avg_confidence > 2 else 'Rendah',
            'evidence': 'Penurunan SSIM, lonjakan aliran optik, perubahan klaster K-Means',
            'interpretation': 'Mengindikasikan potensi penghapusan, penyisipan, atau penyuntingan kasar'
        })
    
    # Check for insertion events
    insertion_count = sum(1 for f in result.frames if f.type == 'anomaly_insertion')
    if insertion_count > 0:
        primary_findings.append({
            'finding': f"Terdeteksi {insertion_count} frame yang berpotensi disisipkan",
            'confidence': 'Tinggi' if avg_confidence > 3 else 'Sedang' if avg_confidence > 2 else 'Rendah',
            'evidence': 'Tidak ada di baseline, anomali ELA, fitur tidak konsisten',
            'interpretation': 'Mengindikasikan konten yang mungkin telah ditambahkan ke video asli'
        })
    
    # Determine overall reliability assessment
    reliability_factors = []
    
    # Factor 1: Multi-method confirmation
    avg_methods = evidence_strength['multi_method_confirmation']['average_methods_per_anomaly']
    reliability_factors.append({
        'factor': 'Konfirmasi Multi-Metode',
        'assessment': f"Rata-rata {avg_methods:.1f} metode mengkonfirmasi setiap anomali",
        'impact': 'Positif' if avg_methods >= 1 else 'Negatif'
    })
    
    # Factor 2: False positive risk
    fp_risk = evidence_strength['false_positive_assessment']['weighted_risk']
    reliability_factors.append({
        'factor': 'Risiko Positif Palsu',
        'assessment': f"Estimasi risiko positif palsu {fp_risk*100:.1f}%",
        'impact': 'Positif' if fp_risk < 0.3 else 'Netral' if fp_risk < 0.5 else 'Negatif'
    })
    
    # Factor 3: Temporal distribution
    temp_dist = anomaly_char['temporal_distribution']['distribution_pattern']
    reliability_factors.append({
        'factor': 'Distribusi Temporal',
        'assessment': f"Anomali menunjukkan pola distribusi {temp_dist}",
        'impact': 'Positif' if temp_dist == 'terkelompok' else 'Netral' if temp_dist == 'sistematis' else 'Positif'
    })
    
    # Factor 4: Technical severity
    avg_severity = anomaly_char['technical_severity']['overall_mean_severity']
    reliability_factors.append({
        'factor': 'Tingkat Keparahan Teknis',
        'assessment': f"Rata-rata keparahan anomali: {avg_severity:.2f} (skala 0-1)",
        'impact': 'Positif' if avg_severity > 0.5 else 'Netral' if avg_severity > 0.3 else 'Negatif'
    })
    
    # Factor 5: Alternative explanations
    most_likely_alt = causality['alternative_explanations'].get('most_likely_alternative', None)
    if most_likely_alt:
        alt_likelihood = causality['alternative_explanations']['all_alternatives'][most_likely_alt]['likelihood']
        reliability_factors.append({
            'factor': 'Penjelasan Alternatif',
            'assessment': f"{most_likely_alt} adalah alternatif dengan kemungkinan {alt_likelihood}",
            'impact': 'Negatif' if alt_likelihood == 'Tinggi' else 'Netral' if alt_likelihood == 'Sedang-Tinggi' else 'Positif'
        })
    
    # Calculate positive vs negative factors
    positive_count = sum(1 for f in reliability_factors if f['impact'] == 'Positif')
    negative_count = sum(1 for f in reliability_factors if f['impact'] == 'Negatif')
    
    # Generate overall reliability statement
    if positive_count >= 3 and negative_count <= 1:
        reliability = "Reliabilitas Tinggi: Bukti sangat kuat mendukung adanya manipulasi video"
    elif positive_count >= 2 and negative_count <= 2:
        reliability = "Reliabilitas Sedang: Bukti menunjukkan kemungkinan adanya manipulasi video"
    elif positive_count >= negative_count:
        reliability = "Reliabilitas Terbatas: Bukti mengindikasikan kemungkinan manipulasi video"
    else:
        reliability = "Reliabilitas Rendah: Bukti tidak meyakinkan atau rentan terhadap penjelasan alternatif"
    
    # Generate recommended actions
    recommended_actions = []
    
    # Action 1: Always recommend based on specific findings
    if primary_findings:
        recommended_actions.append("Investigasi lebih lanjut terhadap segmen anomali spesifik yang diidentifikasi dalam analisis ini")
    
    # Action 2: Based on false positive risk
    if fp_risk > 0.3:
        recommended_actions.append("Dapatkan materi sumber berkualitas lebih tinggi jika memungkinkan untuk mengurangi artefak kompresi")
    
    # Action 3: Based on reliability assessment
    if 'Rendah' in reliability or 'Terbatas' in reliability:
        recommended_actions.append("Terapkan metode forensik tambahan di luar yang digunakan dalam analisis ini")
    
    # Action 4: When alternative explanations are strong
    if most_likely_alt and alt_likelihood in ['Tinggi', 'Sedang-Tinggi']:
        recommended_actions.append(f"Periksa kondisi rekaman asli untuk menyingkirkan {most_likely_alt} sebagai penjelasan")
    
    # Action 5: When dealing with duplications
    if duplication_count > 0:
        recommended_actions.append("Bandingkan segmen yang diduplikasi dengan konteks sekitarnya untuk menentukan tujuan manipulasi")
    
    return {
        'primary_findings': primary_findings,
        'reliability_assessment': reliability,
        'reliability_factors': reliability_factors,
        'recommended_actions': recommended_actions
    }

def create_ferm_visualizations(result: AnalysisResult, ferm: dict, out_dir: Path) -> dict:
    """
    Membuat visualisasi untuk Forensic Evidence Reliability Matrix.
    
    Returns:
        dict: Path ke file visualisasi yang dihasilkan
    """
    visualization_paths = {}
    
    # 1. Evidence Strength Heatmap
    viz_path = create_evidence_strength_heatmap(result, ferm, out_dir)
    if viz_path:
        visualization_paths['evidence_strength'] = str(viz_path)
    
    # 2. Method Correlation Network
    viz_path = create_method_correlation_network(result, ferm, out_dir)
    if viz_path:
        visualization_paths['method_correlation'] = str(viz_path)
    
    # 3. Reliability Factors Assessment
    viz_path = create_reliability_assessment(result, ferm, out_dir)
    if viz_path:
        visualization_paths['reliability_assessment'] = str(viz_path)
    
    # 4. Findings Summary
    viz_path = create_findings_summary(result, ferm, out_dir)
    if viz_path:
        visualization_paths['findings_summary'] = str(viz_path)
    
    return visualization_paths

def create_evidence_strength_heatmap(result: AnalysisResult, ferm: dict, out_dir: Path) -> Path:
    """
    Membuat heatmap yang menunjukkan kekuatan bukti untuk berbagai jenis anomali.
    MODIFIKASI: Judul dan label diterjemahkan ke Bahasa Indonesia.
    """
    # Create figure
    fig, ax = plt.subplots(figsize=(12, 8))
    
    # Prepare data
    anomaly_types = ['Duplikasi', 'Diskontinuitas', 'Penyisipan']
    evidence_methods = ['K-means', 'SSIM', 'Optical Flow', 'ELA', 'SIFT']
    
    # This would be calculated from detailed analysis of which methods detected which anomalies
    # For demonstration, we'll create a simulated heatmap
    data = np.zeros((len(anomaly_types), len(evidence_methods)))
    
    # Count how many times each method confirmed each anomaly type
    for f in result.frames:
        if not f.type.startswith('anomaly'):
            continue
            
        anomaly_idx = -1
        if f.type == 'anomaly_duplication':
            anomaly_idx = 0
        elif f.type == 'anomaly_discontinuity':
            anomaly_idx = 1
        elif f.type == 'anomaly_insertion':
            anomaly_idx = 2
            
        if anomaly_idx == -1:
            continue
            
        if f.evidence_obj.reasons:
            reasons = f.evidence_obj.reasons.split(', ') if isinstance(f.evidence_obj.reasons, str) else f.evidence_obj.reasons
            for reason in reasons:
                if "Adegan" in reason or "K-Means" in reason:
                    data[anomaly_idx, 0] += 1
                if "SSIM" in reason:
                    data[anomaly_idx, 1] += 1
                if "Aliran Optik" in reason:
                    data[anomaly_idx, 2] += 1
                if "ELA" in reason:
                    data[anomaly_idx, 3] += 1
                if "SIFT" in reason or "duplikasi" in reason.lower():
                    data[anomaly_idx, 4] += 1
    
    # Normalize data
    row_sums = data.sum(axis=1, keepdims=True)
    normalized_data = np.zeros_like(data)
    for i in range(data.shape[0]):
        if row_sums[i, 0] > 0:
            normalized_data[i, :] = data[i, :] / row_sums[i, 0]
    
    # Create heatmap
    im = ax.imshow(normalized_data, cmap='YlOrRd')
    
    # Add labels
    ax.set_xticks(np.arange(len(evidence_methods)))
    ax.set_yticks(np.arange(len(anomaly_types)))
    ax.set_xticklabels(evidence_methods)
    ax.set_yticklabels(anomaly_types)
    
    # Rotate x-axis labels
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")
    
    # Add colorbar
    cbar = ax.figure.colorbar(im, ax=ax)
    cbar.ax.set_ylabel("Kekuatan Deteksi Ternormalisasi", rotation=-90, va="bottom")
    
    # Add title and labels
    ax.set_title("Kekuatan Bukti Berdasarkan Jenis Anomali dan Metode Deteksi")
    
    # Add text annotations
    for i in range(len(anomaly_types)):
        for j in range(len(evidence_methods)):
            text = ax.text(j, i, f"{normalized_data[i, j]:.2f}",
                           ha="center", va="center", color="black" if normalized_data[i, j] < 0.7 else "white")
    
    fig.tight_layout()
    
    # Save the visualization
    out_path = out_dir / f"ferm_evidence_strength_{Path(result.video_path).stem}.png"
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    return out_path

def create_method_correlation_network(result: AnalysisResult, ferm: dict, out_dir: Path) -> Path:
    """
    Membuat visualisasi jaringan yang menunjukkan korelasi antar metode deteksi.
    """
    try:
        # This requires networkx which might not be available
        import networkx as nx
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Create graph
        G = nx.Graph()
        
        # Add nodes for detection methods
        methods = ['K-means', 'SSIM', 'Optical Flow', 'ELA', 'SIFT']
        method_colors = {'K-means': 'skyblue', 'SSIM': 'lightgreen', 
                         'Optical Flow': 'salmon', 'ELA': 'purple', 'SIFT': 'orange'}
        
        for method in methods:
            G.add_node(method, color=method_colors[method])
        
        # Calculate edge weights based on how often methods agree
        method_agreements = np.zeros((len(methods), len(methods)))
        
        # Count how many times each pair of methods both detected an anomaly
        for f in result.frames:
            if not f.type.startswith('anomaly') or not f.evidence_obj.reasons:
                continue
                
            detected_methods = set()
            reasons = f.evidence_obj.reasons.split(', ') if isinstance(f.evidence_obj.reasons, str) else f.evidence_obj.reasons
            
            for reason in reasons:
                if "Adegan" in reason or "K-Means" in reason:
                    detected_methods.add('K-means')
                if "SSIM" in reason:
                    detected_methods.add('SSIM')
                if "Aliran Optik" in reason:
                    detected_methods.add('Optical Flow')
                if "ELA" in reason:
                    detected_methods.add('ELA')
                if "SIFT" in reason or "duplikasi" in reason.lower():
                    detected_methods.add('SIFT')
            
            # Add to agreement matrix for every pair of methods
            for i, m1 in enumerate(methods):
                for j, m2 in enumerate(methods):
                    if i != j and m1 in detected_methods and m2 in detected_methods:
                        method_agreements[i, j] += 1
        
        # Add edges based on agreement counts
        for i, m1 in enumerate(methods):
            for j, m2 in enumerate(methods):
                if i < j and method_agreements[i, j] > 0:
                    # Edge weight is proportional to number of agreements
                    G.add_edge(m1, m2, weight=method_agreements[i, j])
        
        # Get position layout
        pos = nx.spring_layout(G, seed=42)
        
        # Get node colors
        node_colors = [G.nodes[node]['color'] for node in G.nodes]
        
        # Get edge weights for width scaling
        edge_weights = [G[u][v]['weight'] for u, v in G.edges]
        max_weight = max(edge_weights) if edge_weights else 1
        edge_widths = [1 + 3 * (w / max_weight) for w in edge_weights]
        
        # Draw the graph
        nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=500, alpha=0.8)
        nx.draw_networkx_edges(G, pos, width=edge_widths, alpha=0.5, edge_color='gray')
        nx.draw_networkx_labels(G, pos, font_size=10, font_family='sans-serif')
        
        # Add edge labels
        edge_labels = {(u, v): f"{G[u][v]['weight']:.0f}" for u, v in G.edges}
        nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=8)
        
        # Add title
        plt.title("Jaringan Korelasi Metode: Seberapa Sering Metode Deteksi Sepakat")
        ax.axis('off')
        
        # Save the visualization
        out_path = out_dir / f"ferm_method_correlation_{Path(result.video_path).stem}.png"
        plt.savefig(out_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return out_path
    except ImportError:
        # Fallback if networkx is not available
        fig, ax = plt.subplots(figsize=(10, 8))
        ax.text(0.5, 0.5, "Jaringan Korelasi Metode\n(membutuhkan pustaka networkx)", 
                ha='center', va='center', fontsize=14)
        ax.axis('off')
        
        out_path = out_dir / f"ferm_method_correlation_{Path(result.video_path).stem}.png"
        plt.savefig(out_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return out_path

# GANTI FUNGSI INI DI FILE ForensikVideo.py

def create_reliability_assessment(result: AnalysisResult, ferm: dict, out_dir: Path) -> Path:
    """
    Membuat visualisasi untuk penilaian reliabilitas bukti forensik.
    MODIFIKASI: Judul dan label diterjemahkan ke Bahasa Indonesia.
    FIX: Memperbaiki masalah teks tumpang tindih.
    """
    # Get reliability factors
    reliability_factors = ferm['conclusion'].get('reliability_factors', [])
    
    if not reliability_factors:
        # Create empty placeholder
        fig, ax = plt.subplots(figsize=(12, 7)) # Tetap gunakan ukuran yang baik
        ax.text(0.5, 0.5, "Faktor reliabilitas tidak tersedia", ha='center', va='center', fontsize=14)
        ax.axis('off')
    else:
        # === SOLUSI 1: Tingkatkan ukuran vertikal gambar agar ada lebih banyak ruang ===
        # Jumlah faktor menentukan tinggi gambar, min 7, maks 14 inci
        num_factors = len(reliability_factors)
        fig_height = max(7, min(14, num_factors * 1.8))
        fig, ax = plt.subplots(figsize=(12, fig_height))
    
        # Extract factor names and impact
        factors = [f['factor'] for f in reliability_factors]
        impacts = [f['impact'] for f in reliability_factors]
        assessments = [f['assessment'] for f in reliability_factors]
        
        # Convert impacts to numeric values
        impact_values = []
        for impact in impacts:
            if impact == 'Positif':
                impact_values.append(1)
            elif impact == 'Netral':
                impact_values.append(0)
            else:  # Negative
                impact_values.append(-1)
        
        # Define colors based on impact
        colors = ['#28a745' if i == 'Positif' else '#ffc107' if i == 'Netral' else '#dc3545' for i in impacts]
        
        # Create horizontal bar chart
        bars = ax.barh(factors, impact_values, color=colors, alpha=0.7, height=0.8)
        
        # Add assessment text, wrap long text
        import textwrap
        for i, assessment in enumerate(assessments):
            # Bungkus teks agar tidak terlalu panjang
            wrapped_assessment = '\n'.join(textwrap.wrap(assessment, width=30))
            if impact_values[i] >= 0:
                ax.text(1.05, i, wrapped_assessment, va='center', ha='left', fontsize=9)
            else:
                ax.text(-1.05, i, wrapped_assessment, va='center', ha='right', fontsize=9)
        
        # Add labels and title
        ax.set_xlim(-1.6, 1.6)
        ax.set_xticks([-1, 0, 1])
        ax.set_xticklabels(['Dampak Negatif', 'Netral', 'Dampak Positif'])
        ax.axvline(x=0, color='black', linestyle='-', alpha=0.3)
        
        # Add reliability assessment as title
        reliability = ferm['conclusion'].get('reliability_assessment', 'Penilaian reliabilitas tidak tersedia')
        plt.title(f"Penilaian Reliabilitas Bukti\n{reliability}", fontsize=14, wrap=True)
        
        # Add grid
        ax.grid(True, axis='x', alpha=0.3)

    # === SOLUSI 2: Gunakan tight_layout untuk menyesuaikan plot secara otomatis ===
    plt.tight_layout(pad=3.0) # pad=3.0 memberikan lebih banyak bantalan
    plt.subplots_adjust(left=0.25, right=0.95)
    
    # Save the visualization
    out_path = out_dir / f"ferm_reliability_{Path(result.video_path).stem}.png"
    # === SOLUSI 3: Tambahkan bbox_inches='tight' untuk memotong ruang putih ekstra ===
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    return out_path
    
    # Create figure
    fig, ax = plt.subplots(figsize=(12, 7))
    
    # Extract factor names and impact
    factors = [f['factor'] for f in reliability_factors]
    impacts = [f['impact'] for f in reliability_factors]
    assessments = [f['assessment'] for f in reliability_factors]
    
    # Convert impacts to numeric values
    impact_values = []
    for impact in impacts:
        if impact == 'Positif':
            impact_values.append(1)
        elif impact == 'Netral':
            impact_values.append(0)
        else:  # Negative
            impact_values.append(-1)
    
    # Define colors based on impact
    colors = ['green' if i == 'Positif' else 'gold' if i == 'Netral' else 'red' for i in impacts]
    
    # Create horizontal bar chart
    bars = ax.barh(factors, impact_values, color=colors, alpha=0.7)
    
    # Add assessment text
    for i, assessment in enumerate(assessments):
        if impact_values[i] >= 0:
            ax.text(1.1, i, assessment, va='center', fontsize=9)
        else:
            ax.text(-1.1, i, assessment, va='center', ha='right', fontsize=9)
    
    # Add labels and title
    ax.set_xlim(-1.5, 1.5)
    ax.set_xticks([-1, 0, 1])
    ax.set_xticklabels(['Dampak Negatif', 'Netral', 'Dampak Positif'])
    ax.axvline(x=0, color='black', linestyle='-', alpha=0.3)
    
    # Add reliability assessment as title
    reliability = ferm['conclusion'].get('reliability_assessment', 'Penilaian reliabilitas tidak tersedia')
    plt.title(f"Penilaian Reliabilitas Bukti\n{reliability}", fontsize=14)
    
    # Add grid
    ax.grid(True, axis='x', alpha=0.3)
    
    # Save the visualization
    out_path = out_dir / f"ferm_reliability_{Path(result.video_path).stem}.png"
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    return out_path

def create_findings_summary(result: AnalysisResult, ferm: dict, out_dir: Path) -> Path:
    """
    Membuat visualisasi ringkasan temuan forensik utama.
    """
    # Get primary findings
    findings = ferm['conclusion'].get('primary_findings', [])
    
    # Create figure with subplots - different layout based on findings count
    n_findings = len(findings)
    if n_findings == 0:
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, "Tidak ada temuan signifikan yang terdeteksi", ha='center', va='center', fontsize=14)
        ax.axis('off')
    elif n_findings == 1:
        fig, axs = plt.subplots(1, 1, figsize=(12, 6))
        axs = [axs]
    elif n_findings == 2:
        fig, axs = plt.subplots(1, 2, figsize=(14, 6))
    else:
        fig, axs = plt.subplots(2, 2, figsize=(14, 10))
        axs = axs.flatten()
    
    # Add title to figure
    if n_findings > 0:
        fig.suptitle("Temuan Kunci Forensik", fontsize=16, fontweight='bold')
        
        # Add findings to subplots
        for i, finding in enumerate(findings[:min(n_findings, 4)]):  # Limit to 4 findings
            ax = axs[i]
            
            # Extract finding details
            title = finding.get('finding', 'Temuan tidak spesifik')
            confidence = finding.get('confidence', 'Tidak diketahui')
            evidence = finding.get('evidence', 'Tidak spesifik')
            interpretation = finding.get('interpretation', 'Tidak ada interpretasi')
            
            # Create colored box based on confidence
            if confidence == 'Tinggi':
                color = 'lightgreen'
            elif confidence == 'Sedang':
                color = 'khaki'
            else:
                color = 'salmon'
                
            # Add content to subplot
            ax.add_patch(plt.Rectangle((0, 0), 1, 1, facecolor=color, alpha=0.3))
            
            # Add text content
            ax.text(0.5, 0.85, title, ha='center', va='center', fontsize=12, fontweight='bold',
                   wrap=True, bbox=dict(facecolor='white', alpha=0.7))
            
            ax.text(0.5, 0.7, f"Kepercayaan: {confidence}", ha='center', va='center', fontsize=11)
            
            ax.text(0.5, 0.5, f"Bukti:\n{evidence}", ha='center', va='center', fontsize=10,
                   wrap=True, bbox=dict(facecolor='white', alpha=0.4))
            
            ax.text(0.5, 0.2, f"Interpretasi:\n{interpretation}", ha='center', va='center', 
                   fontsize=10, fontstyle='italic', wrap=True)
            
            # Remove axes
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
            ax.axis('off')
    
    # Hide any unused subplots
    if n_findings > 0 and n_findings < len(axs):
        for i in range(n_findings, len(axs)):
            axs[i].axis('off')
    
    # Add recommended actions at the bottom
    if n_findings > 0:
        recommended_actions = ferm['conclusion'].get('recommended_actions', [])
        if recommended_actions:
            actions_text = "Rekomendasi Tindakan:\n" + "\n".join([f"• {action}" for action in recommended_actions])
            fig.text(0.5, 0.02, actions_text, ha='center', va='bottom', fontsize=10, 
                    bbox=dict(facecolor='lightyellow', alpha=0.5))
    
    # Save the visualization
    out_path = out_dir / f"ferm_findings_{Path(result.video_path).stem}.png"
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    return out_path

def create_zip_archive(result: AnalysisResult, out_dir: Path) -> Path | None:
    """Membuat arsip ZIP yang berisi seluruh laporan dan artefak analisis."""

    zip_filename = f"Laporan_Lengkap_{Path(result.video_path).stem}.zip"
    zip_path = out_dir / zip_filename

    log(f"  {Icons.INFO} Membuat arsip ZIP lengkap...")

    try:
        # Kumpulkan semua file terlebih dahulu agar file ZIP tidak ikut terarsip
        files_to_zip: list[tuple[Path, Path]] = []
        for root, dirs, files in os.walk(out_dir):
            for file in files:
                file_path = Path(root) / file
                if file_path.resolve() != zip_path.resolve():
                    arcname = file_path.relative_to(out_dir)
                    files_to_zip.append((file_path, arcname))

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path, arcname in files_to_zip:
                zipf.write(file_path, arcname=arcname)

        log(f"  ✅ Arsip ZIP berhasil dibuat: {zip_path.name}")
        return zip_path
    except Exception as e:
        log(f"  {Icons.ERROR} Gagal membuat arsip ZIP: {e}")
        return None


###############################################################################
# PIPELINE 5-TAHAP
###############################################################################

# --- TAHAP 1: PRA-PEMROSESAN & EKSTRAKSI FITUR DASAR ---
def run_tahap_1_pra_pemrosesan(video_path: Path, out_dir: Path, fps: int) -> AnalysisResult | None:
    print_stage_banner(1, "Pra-pemrosesan & Ekstraksi Fitur Dasar", Icons.COLLECTION,
                       "Mengamankan bukti, mengekstrak metadata, menormalisasi frame, dan menerapkan metode K-Means.")

    log(f"  {Icons.IDENTIFICATION} Melakukan preservasi bukti dengan hashing SHA-256...")
    preservation_hash = calculate_sha256(video_path)
    log(f"  ✅ Hash SHA-256: {preservation_hash}")

    log(f"  {Icons.PRESERVATION} Mengekstrak metadata detail dengan FFprobe...")
    metadata_raw = ffprobe_metadata(video_path)
    metadata = parse_ffprobe_output(metadata_raw)

    log(f"  {Icons.COLLECTION} Mengekstrak, menormalisasi, dan membandingkan frame @ {fps} FPS...")
    frames_dir_root = out_dir / f"frames_{video_path.stem}"
    extracted_paths = extract_frames_with_normalization(video_path, frames_dir_root, fps)
    if not extracted_paths:
        log(f"  {Icons.ERROR} Gagal mengekstrak frame. Pastikan video valid dan FFmpeg/OpenCV berfungsi.")
        return None
    log(f"  ✅ {len(extracted_paths)} set frame (original, normalized, comparison) berhasil diekstrak.")

    log(f"  {Icons.EXAMINATION} Menghitung pHash untuk setiap frame (menggunakan frame ternormalisasi)...")
    frames = []
    for idx, (p_orig, p_norm, p_comp) in enumerate(tqdm(extracted_paths, desc="    pHash", leave=False, bar_format='{l_bar}{bar}{r_bar}')):
        try:
            with Image.open(p_norm) as img:
                frame_hash = str(imagehash.average_hash(img))
            frames.append(FrameInfo(
                index=idx,
                timestamp=idx / fps,
                img_path_original=p_orig,
                img_path=p_norm, # img_path utama menunjuk ke versi ternormalisasi
                img_path_comparison=p_comp,
                hash=frame_hash
            ))
        except Exception as e:
            log(f"  {Icons.ERROR} Gagal memproses frame set {idx}: {e}")

    log(f"  {Icons.EXAMINATION} METODE UTAMA: Menganalisis layout warna global (K-Means)...")
    histograms = []
    for f in tqdm(frames, desc="    Histogram (Normalized)", leave=False, bar_format='{l_bar}{bar}{r_bar}'):
        img = cv2.imread(f.img_path) # Baca dari frame ternormalisasi
        if img is None: continue
        hist = cv2.calcHist([img], [0, 1, 2], None, [8, 8, 8], [0, 256, 0, 256, 0, 256])
        cv2.normalize(hist, hist)
        histograms.append(hist.flatten())

    # ====== [NEW] False-Positive Fix June-2025 ======
    if histograms:
        scene_variance = float(np.var(histograms))
        if scene_variance < 0.15:
            CONFIG["KMEANS_CLUSTERS"] = 5
    # ====== [END NEW] ======

    kmeans_artifacts = {}
    if histograms:
        actual_n_clusters = min(CONFIG["KMEANS_CLUSTERS"], len(histograms))
        if actual_n_clusters >= 2:
            kmeans = KMeans(n_clusters=actual_n_clusters, random_state=42, n_init='auto').fit(histograms)
            labels = kmeans.labels_.tolist()
            for f, label in zip(frames, labels):
                f.color_cluster = int(label)
            log(f"  -> Klasterisasi K-Means selesai. {len(frames)} frame dikelompokkan ke dalam {actual_n_clusters} klaster.")

            # --- PEMBUATAN ARTEFAK K-MEANS DETAIL ---
            log(f"  {Icons.ANALYSIS} Membuat artefak visualisasi detail untuk K-Means...")
            kmeans_dir = out_dir / "kmeans_artifacts"
            kmeans_dir.mkdir(exist_ok=True)

            # 1. Plot Distribusi Klaster
            cluster_counts = Counter(labels)
            plt.figure(figsize=(10, 5))
            plt.bar(list(cluster_counts.keys()), list(cluster_counts.values()), color='cornflowerblue')
            plt.title('Distribusi Frame per Klaster K-Means', fontsize=14)
            plt.xlabel('Nomor Klaster', fontsize=12)
            plt.ylabel('Jumlah Frame', fontsize=12)
            plt.xticks(range(actual_n_clusters))
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            dist_path = kmeans_dir / "kmeans_distribution.png"
            plt.savefig(dist_path, bbox_inches="tight"); plt.close()
            kmeans_artifacts['distribution_plot_path'] = str(dist_path)

            # 2. Palet Warna dan Sampel Frame per Klaster
            kmeans_artifacts['clusters'] = []
            for i in range(actual_n_clusters):
                cluster_indices = [idx for idx, label in enumerate(labels) if label == i]
                if not cluster_indices: continue

                # Buat palet warna dari rata-rata frame di klaster
                avg_color_img = np.zeros((100, 400, 3), np.uint8)
                # Ambil satu frame representatif untuk diekstrak warnanya
                sample_frame_path = frames[cluster_indices[0]].img_path
                sample_img = cv2.imread(sample_frame_path)
                if sample_img is not None:
                    pixels = sample_img.reshape(-1, 3).astype(np.float32)
                    palette_kmeans = KMeans(n_clusters=5, random_state=42, n_init='auto').fit(pixels)
                    for j, color in enumerate(palette_kmeans.cluster_centers_):
                        cv2.rectangle(avg_color_img, (j*80, 0), ((j+1)*80, 100), color.astype(int).tolist(), -1)

                palette_path = kmeans_dir / f"cluster_{i}_palette.png"
                cv2.imwrite(str(palette_path), avg_color_img)

                # Buat montase sampel frame
                sample_frames_to_show = [frames[j] for j in cluster_indices[:CONFIG["KMEANS_SAMPLES_PER_CLUSTER"]]]
                montage_h = (Image.open(sample_frames_to_show[0].img_path_original).height if sample_frames_to_show else 180)
                montage_w = (Image.open(sample_frames_to_show[0].img_path_original).width if sample_frames_to_show else 320)
                montage_img = Image.new('RGB', (montage_w * len(sample_frames_to_show), montage_h))
                for k, f_info in enumerate(sample_frames_to_show):
                    with Image.open(f_info.img_path_original) as img:
                        img = img.resize((montage_w, montage_h))
                        montage_img.paste(img, (k * montage_w, 0))

                montage_path = kmeans_dir / f"cluster_{i}_samples.jpg"
                montage_img.save(montage_path)

                kmeans_artifacts['clusters'].append({
                    'id': i,
                    'count': len(cluster_indices),
                    'palette_path': str(palette_path),
                    'samples_montage_path': str(montage_path)
                })
            log(f"  -> Artefak K-Means berhasil dibuat di direktori {kmeans_dir.name}")

    result = AnalysisResult(
        video_path=str(video_path),
        preservation_hash=preservation_hash,
        metadata=metadata,
        metadata_raw=metadata_raw,
        metadata_analysis=analyze_video_metadata(metadata),
        frames=frames,
        kmeans_artifacts=kmeans_artifacts
    )

    log(f"  {Icons.SUCCESS} Tahap 1 Selesai.")
    return result

# --- TAHAP 2: ANALISIS ANOMALI TEMPORAL & KOMPARATIF ---
def run_tahap_2_analisis_temporal(result: AnalysisResult, baseline_result: AnalysisResult | None = None):
    print_stage_banner(2, "Analisis Anomali Temporal & Komparatif", Icons.ANALYSIS,
                       "Menganalisis aliran optik, SSIM, dan perbandingan dengan baseline jika ada.")
    frames = result.frames
    prev_gray = None

    log(f"  {Icons.EXAMINATION} Menghitung Aliran Optik & SSIM antar frame (menggunakan frame ternormalisasi)...")
    for f_idx, f in enumerate(tqdm(frames, desc="    Temporal", leave=False, bar_format='{l_bar}{bar}{r_bar}')):
        current_gray = cv2.imread(f.img_path, cv2.IMREAD_GRAYSCALE) # f.img_path adalah frame ternormalisasi
        if current_gray is not None:
            if prev_gray is not None and prev_gray.shape == current_gray.shape:
                data_range = float(current_gray.max() - current_gray.min())
                if data_range > 0:
                    ssim_score = ssim(prev_gray, current_gray, data_range=data_range)
                    f.ssim_to_prev = float(ssim_score)
                else:
                    f.ssim_to_prev = 1.0 # Frames are identical if data_range is 0

                if current_gray.dtype == prev_gray.dtype:
                    try:
                        flow = cv2.calcOpticalFlowFarneback(prev_gray, current_gray, None, 0.5, 3, 15, 3, 5, 1.2, 0)
                        mag, _ = cv2.cartToPolar(flow[..., 0], flow[..., 1])
                        f.optical_flow_mag = float(np.mean(mag))
                    except cv2.error as e:
                        log(f"  {Icons.ERROR} OpenCV error during optical flow for frame {f.index}: {e}")
                        f.optical_flow_mag = 0.0
                else:
                    f.optical_flow_mag = 0.0
            else:
                f.ssim_to_prev = 1.0
                f.optical_flow_mag = 0.0

        prev_gray = current_gray

    if baseline_result:
        log(f"  {Icons.ANALYSIS} Melakukan analisis komparatif terhadap video baseline...")
        base_hashes = {bf.hash for bf in baseline_result.frames if bf.hash}
        insertion_count = 0
        for f_sus in frames:
            if f_sus.hash and f_sus.hash not in base_hashes:
                f_sus.type = "anomaly_insertion"
                f_sus.evidence_obj.reasons.append("Frame tidak ada di baseline")
                f_sus.evidence_obj.confidence = "SANGAT TINGGI"
                insertion_count += 1
        log(f"  -> Terdeteksi {insertion_count} frame sisipan potensial.")

    log(f"  {Icons.SUCCESS} Tahap 2 Selesai.")

# --- TAHAP 3: SINTESIS BUKTI & INVESTIGASI MENDALAM ---
def run_tahap_3_sintesis_bukti(result: AnalysisResult, out_dir: Path):
    print_stage_banner(3, "Sintesis Bukti & Investigasi Mendalam", "🔬",
                       "Mengkorelasikan semua temuan dan melakukan analisis ELA/SIFT pada anomali terkuat dengan penjelasan detail.")
    frames = result.frames
    n = len(frames)
    if n < 2: return

    # Inisialisasi struktur untuk analisis detail
    result.detailed_anomaly_analysis = {
        'temporal_discontinuities': [],
        'duplication_analysis': [],
        'compression_anomalies': [],
        'statistical_outliers': []
    }

    log(f"  {Icons.ANALYSIS} ANALISIS 1: Deteksi Diskontinuitas Temporal...")
    log(f"  📖 Penjelasan: Diskontinuitas temporal adalah perubahan mendadak antara frame yang berurutan.")
    log(f"     Ini bisa mengindikasikan penghapusan frame, penyisipan konten, atau editing yang kasar.")

    # Kalkulasi metrik tambahan untuk setiap frame
    log(f"  {Icons.EXAMINATION} Menghitung metrik detail untuk setiap frame...")
    for f in tqdm(frames, desc="    Metrik Frame", leave=False):
        metrics = calculate_frame_metrics(f.img_path_original)
        f.edge_density = metrics.get('edge_density')
        f.blur_metric = metrics.get('blur_metric')
        f.evidence_obj.detailed_analysis['frame_metrics'] = metrics

    # Analisis metrik diskontinuitas dengan penjelasan
    flow_mags = [f.optical_flow_mag for f in frames if f.optical_flow_mag is not None]

    if flow_mags:
        # Gunakan metode statistik yang lebih robust
        filtered_flow_mags = [m for m in flow_mags if m > 0.0]
        if len(filtered_flow_mags) > 1:
            median_flow = np.median(filtered_flow_mags)
            mad_flow = stats.median_abs_deviation(filtered_flow_mags)
            mad_flow = 1e-9 if mad_flow == 0 else mad_flow

            # Hitung persentil untuk context
            p25 = np.percentile(filtered_flow_mags, 25)
            p75 = np.percentile(filtered_flow_mags, 75)
            p95 = np.percentile(filtered_flow_mags, 95)

            log(f"  📊 Statistik Aliran Optik:")
            log(f"     - Median: {median_flow:.3f}")
            log(f"     - MAD (Median Absolute Deviation): {mad_flow:.3f}")

            # Deteksi anomali dengan Z-score
            for f in frames:
                if f.optical_flow_mag is not None and f.optical_flow_mag > 0:
                    if mad_flow != 0:
                        z_score = 0.6745 * (f.optical_flow_mag - median_flow) / mad_flow
                        if abs(z_score) > CONFIG["OPTICAL_FLOW_Z_THRESH"]:
                            f.evidence_obj.reasons.append("Lonjakan Aliran Optik")
                            f.evidence_obj.metrics["optical_flow_z_score"] = round(z_score, 2)

                            # Tambahkan penjelasan detail
                            explanation = {
                                "type": "optical_flow_spike",
                                "frame_index": f.index,
                                "timestamp": f.timestamp,
                                "severity": "high" if abs(z_score) > 6 else "medium",
                                "technical_explanation": (
                                    f"Frame ini menunjukkan pergerakan piksel yang {abs(z_score):.1f}x "
                                    "lebih besar dari normal."
                                ),
                                "simple_explanation": (
                                    "Terjadi perubahan gambar yang sangat mendadak, "
                                    "seperti perpindahan kamera yang kasar atau cut yang tidak halus."
                                ),
                                "metrics": {
                                    "flow_magnitude": f.optical_flow_mag,
                                    "z_score": z_score,
                                    "median_flow": median_flow,
                                    "deviation_percentage": (
                                        (f.optical_flow_mag - median_flow) / median_flow * 100
                                    )
                                    if median_flow > 0
                                    else 0,
                                },
                            }
                            f.evidence_obj.explanations['optical_flow'] = explanation
                            result.detailed_anomaly_analysis['temporal_discontinuities'].append(explanation)

    # Analisis SSIM dengan konteks yang lebih kaya
    log(f"\n  {Icons.ANALYSIS} ANALISIS 2: Deteksi Penurunan Kemiripan Struktural (SSIM)...")
    log(f"  📖 Penjelasan: SSIM mengukur seberapa mirip dua gambar secara struktural.")
    log(f"     Nilai 1.0 = identik, nilai < 0.7 = sangat berbeda. Penurunan drastis = kemungkinan manipulasi.")

    ssim_values = [f.ssim_to_prev for f in frames if f.ssim_to_prev is not None]
    if ssim_values:
        ssim_mean = np.mean(ssim_values)
        ssim_std = np.std(ssim_values)
        log(f"  📊 Statistik SSIM: Mean={ssim_mean:.3f}, Std={ssim_std:.3f}")

    for i in range(1, n):
        f_curr, f_prev = frames[i], frames[i - 1]
        if f_curr.ssim_to_prev is not None and f_prev.ssim_to_prev is not None:
            ssim_drop = f_prev.ssim_to_prev - f_curr.ssim_to_prev

            # Deteksi penurunan drastis
            if ssim_drop > CONFIG["SSIM_DISCONTINUITY_DROP"]:
                f_curr.evidence_obj.reasons.append("Penurunan Drastis SSIM")
                f_curr.evidence_obj.metrics["ssim_drop"] = round(ssim_drop, 4)

                explanation = {
                    'type': 'ssim_drop',
                    'frame_index': f_curr.index,
                    'timestamp': f_curr.timestamp,
                    'severity': 'high' if ssim_drop > 0.5 else 'medium',
                    'technical_explanation': f"SSIM turun {ssim_drop:.3f} dari frame sebelumnya ({f_prev.ssim_to_prev:.3f} → {f_curr.ssim_to_prev:.3f}).",
                    'simple_explanation': "Frame ini sangat berbeda dari frame sebelumnya, mungkin ada potongan atau sisipan.",
                    'metrics': {
                        'ssim_current': f_curr.ssim_to_prev,
                        'ssim_previous': f_prev.ssim_to_prev,
                        'drop_amount': ssim_drop,
                        'drop_percentage': (ssim_drop / f_prev.ssim_to_prev * 100) if f_prev.ssim_to_prev > 0 else 0
                    }
                }
                f_curr.evidence_obj.explanations['ssim_drop'] = explanation
                result.detailed_anomaly_analysis['temporal_discontinuities'].append(explanation)

            # Deteksi nilai SSIM sangat rendah
            elif f_curr.ssim_to_prev < 0.7:
                f_curr.evidence_obj.reasons.append("SSIM Sangat Rendah")
                f_curr.evidence_obj.metrics["ssim_absolute_low"] = round(f_curr.ssim_to_prev, 4)

                explanation = {
                    'type': 'ssim_low',
                    'frame_index': f_curr.index,
                    'timestamp': f_curr.timestamp,
                    'severity': 'medium',
                    'technical_explanation': f"SSIM sangat rendah ({f_curr.ssim_to_prev:.3f}), menunjukkan perbedaan struktural yang signifikan.",
                    'simple_explanation': "Frame ini memiliki struktur visual yang sangat berbeda dari frame sebelumnya.",
                    'metrics': {
                        'ssim_value': f_curr.ssim_to_prev,
                        'threshold': 0.7,
                        'below_threshold_by': 0.7 - f_curr.ssim_to_prev
                    }
                }
                f_curr.evidence_obj.explanations['ssim_low'] = explanation

    # Analisis perubahan klaster warna dengan konteks
    log(f"\n  {Icons.ANALYSIS} ANALISIS 3: Deteksi Perubahan Adegan (K-Means)...")
    log(f"  📖 Penjelasan: K-Means mengelompokkan frame berdasarkan palet warna dominan.")
    log(f"     Perubahan klaster = perubahan adegan. Perubahan yang terlalu sering = kemungkinan editing.")

    scene_changes = []
    for i in range(1, n):
        f_curr, f_prev = frames[i], frames[i - 1]
        if f_curr.color_cluster is not None and f_prev.color_cluster is not None and f_curr.color_cluster != f_prev.color_cluster:
            f_curr.evidence_obj.reasons.append("Perubahan Adegan (dari K-Means)")
            f_curr.evidence_obj.metrics["color_cluster_jump"] = f"{f_prev.color_cluster} → {f_curr.color_cluster}"

            scene_change = {
                'frame_index': f_curr.index,
                'timestamp': f_curr.timestamp,
                'from_cluster': f_prev.color_cluster,
                'to_cluster': f_curr.color_cluster,
                'time_since_last_change': 0  # Will be calculated
            }
            scene_changes.append(scene_change)

            explanation = {
                'type': 'scene_change',
                'frame_index': f_curr.index,
                'timestamp': f_curr.timestamp,
                'technical_explanation': f"Perubahan dari klaster warna {f_prev.color_cluster} ke {f_curr.color_cluster}.",
                'simple_explanation': "Terjadi perubahan adegan atau sudut pandang kamera.",
                'metrics': {
                    'from_cluster': f_prev.color_cluster,
                    'to_cluster': f_curr.color_cluster
                }
            }
            f_curr.evidence_obj.explanations['scene_change'] = explanation

    # Hitung frekuensi perubahan adegan
    if scene_changes:
        for i in range(1, len(scene_changes)):
            scene_changes[i]['time_since_last_change'] = scene_changes[i]['timestamp'] - scene_changes[i-1]['timestamp']

        avg_scene_duration = np.mean([sc['time_since_last_change'] for sc in scene_changes[1:]]) if len(scene_changes) > 1 else 0
        log(f"  📊 Total perubahan adegan: {len(scene_changes)}")
        log(f"     Durasi rata-rata per adegan: {avg_scene_duration:.2f} detik")

    # METODE PENDUKUNG: Verifikasi duplikasi dengan analisis mendalam
    log(f"\n  {Icons.EXAMINATION} METODE PENDUKUNG 1: Analisis Duplikasi Frame (SIFT+RANSAC)...")
    log(f"  📖 Penjelasan: SIFT mendeteksi titik-titik unik dalam gambar. Jika dua frame memiliki")
    log(f"     banyak titik yang cocok sempurna, kemungkinan besar frame tersebut diduplikasi.")

    hash_map = defaultdict(list)
    for f in frames:
        if f.hash: hash_map[f.hash].append(f.index)

    dup_candidates = {k: v for k, v in hash_map.items() if len(v) > 1}

    if dup_candidates:
        log(f"  🔍 Ditemukan {len(dup_candidates)} grup kandidat duplikasi untuk diverifikasi...")

        # Loop through duplicate candidates
        for hash_val, indices in dup_candidates.items():
            for i in range(1, len(indices)):
                idx1, idx2 = indices[0], indices[i]
                p1 = Path(frames[idx1].img_path_original)
                p2 = Path(frames[idx2].img_path_original)

                # Cek SSIM terlebih dahulu
                im1 = cv2.imread(str(p1), cv2.IMREAD_GRAYSCALE)
                im2 = cv2.imread(str(p2), cv2.IMREAD_GRAYSCALE)
                if im1 is None or im2 is None: continue
                if im1.shape != im2.shape: continue

                data_range = float(im1.max() - im1.min())
                if data_range == 0: continue
                ssim_val = ssim(im1, im2, data_range=data_range)

                if ssim_val > CONFIG["DUPLICATION_SSIM_CONFIRM"]:
                    # Analisis SIFT detail
                    sift_result = compare_sift_enhanced(p1, p2, out_dir)

                    if sift_result.get('success') and sift_result.get('inliers', 0) >= CONFIG["SIFT_MIN_MATCH_COUNT"]:
                        f_dup = frames[idx2]
                        f_dup.type = "anomaly_duplication"
                        f_dup.evidence_obj.reasons.append(f"Duplikasi dari frame {idx1}")
                        f_dup.evidence_obj.metrics.update({
                            "source_frame": idx1,
                            "ssim_to_source": round(ssim_val, 4),
                            "sift_inliers": sift_result['inliers'],
                            "sift_good_matches": sift_result['good_matches'],
                            "sift_inlier_ratio": round(sift_result['inlier_ratio'], 3)
                        })

                        if sift_result.get('visualization_path'):
                            f_dup.evidence_obj.sift_path = sift_result['visualization_path']
                            f_dup.evidence_obj.visualizations['sift_matches'] = sift_result['visualization_path']

                        if sift_result.get('heatmap_path'):
                            f_dup.evidence_obj.visualizations['sift_heatmap'] = sift_result['heatmap_path']

                        # Penjelasan detail duplikasi
                        duplication_analysis = {
                            'type': 'frame_duplication',
                            'duplicate_frame': idx2,
                            'source_frame': idx1,
                            'timestamp_duplicate': frames[idx2].timestamp,
                            'timestamp_source': frames[idx1].timestamp,
                            'time_gap': frames[idx2].timestamp - frames[idx1].timestamp,
                            'confidence': 'very_high' if sift_result['inlier_ratio'] > 0.8 else 'high',
                            'technical_explanation': f"Frame {idx2} adalah duplikasi dari frame {idx1} dengan {sift_result['inliers']} titik fitur yang cocok sempurna ({sift_result['inlier_ratio']:.1%} akurasi).",
                            'simple_explanation': f"Frame pada detik {frames[idx2].timestamp:.2f} adalah salinan persis dari frame pada detik {frames[idx1].timestamp:.2f}. Ini sering digunakan untuk memperpanjang durasi video atau menyembunyikan penghapusan konten.",
                            'sift_analysis': sift_result,
                            'implications': "Duplikasi frame dapat mengindikasikan: (1) Usaha memperpanjang durasi, (2) Menutupi frame yang dihapus, (3) Teknik editing untuk transisi"
                        }
                        f_dup.evidence_obj.explanations['duplication'] = duplication_analysis
                        result.detailed_anomaly_analysis['duplication_analysis'].append(duplication_analysis)

    # METODE PENDUKUNG: ELA dengan analisis regional
    log(f"\n  {Icons.ANALYSIS} METODE PENDUKUNG 2: Error Level Analysis (ELA) untuk Anomali Signifikan...")
    log(f"  📖 Penjelasan: ELA mendeteksi area yang telah diedit dengan melihat perbedaan kompresi.")
    log(f"     Area yang lebih terang dalam ELA = kemungkinan telah dimodifikasi atau disisipkan.")

    # Buat direktori untuk visualisasi tambahan
    detail_viz_dir = out_dir / "detailed_visualizations"
    detail_viz_dir.mkdir(exist_ok=True)

    for f in tqdm(frames, desc="    Analisis ELA & Sintesis", leave=False):
        # First, ensure reasons is a list
        if isinstance(f.evidence_obj.reasons, str):
            f.evidence_obj.reasons = [r.strip() for r in f.evidence_obj.reasons.split(',')]

        if f.evidence_obj.reasons:
            if f.type == "original":
                f.type = "anomaly_discontinuity"

            # Tentukan tingkat kepercayaan berdasarkan jumlah bukti
            num_reasons = len(f.evidence_obj.reasons)
            if f.type == "anomaly_duplication" or f.type == "anomaly_insertion":
                f.evidence_obj.confidence = "SANGAT TINGGI"
            elif num_reasons > 2:
                f.evidence_obj.confidence = "TINGGI"
            elif num_reasons > 1:
                f.evidence_obj.confidence = "SEDANG"
            else:
                f.evidence_obj.confidence = "RENDAH"

            # Lakukan ELA untuk anomali dengan kepercayaan sedang ke atas
            if f.evidence_obj.confidence in ["SEDANG", "TINGGI", "SANGAT TINGGI"] and f.type not in ["anomaly_duplication", "anomaly_insertion"]:
                ela_result = perform_ela(Path(f.img_path_original))
                if ela_result:
                    ela_path, max_diff, ela_array = ela_result
                    f.evidence_obj.ela_path = str(ela_path)

                    # Analisis regional ELA
                    regional_analysis = analyze_ela_regions(ela_array)

                    if regional_analysis['suspicious_count'] > 0:
                        if "Anomali Kompresi (ELA)" not in f.evidence_obj.reasons:
                            f.evidence_obj.reasons.append("Anomali Kompresi (ELA)")
                        f.evidence_obj.metrics["ela_max_difference"] = max_diff
                        f.evidence_obj.metrics["ela_suspicious_regions"] = regional_analysis['suspicious_count']

                        # Upgrade confidence jika ditemukan area mencurigakan
                        if regional_analysis['suspicious_count'] > 5:
                            if f.evidence_obj.confidence == "SEDANG":
                                f.evidence_obj.confidence = "TINGGI"
                            elif f.evidence_obj.confidence == "TINGGI":
                                f.evidence_obj.confidence = "SANGAT TINGGI"

                        # Buat visualisasi ELA dengan highlight area mencurigakan
                        ela_viz_path = create_ela_visualization(
                            Path(f.img_path_original),
                            ela_array,
                            regional_analysis,
                            detail_viz_dir
                        )
                        if ela_viz_path:
                            f.evidence_obj.visualizations['ela_detailed'] = str(ela_viz_path)

                        # Penjelasan detail ELA
                        ela_explanation = {
                            'type': 'compression_anomaly',
                            'frame_index': f.index,
                            'timestamp': f.timestamp,
                            'max_difference': max_diff,
                            'suspicious_regions': regional_analysis['suspicious_regions'][:5],  # Top 5
                            'total_suspicious_areas': regional_analysis['suspicious_count'],
                            'technical_explanation': f"ELA menunjukkan {regional_analysis['suspicious_count']} area dengan perbedaan kompresi tinggi (max: {max_diff}). Area ini kemungkinan telah diedit atau disisipkan.",
                            'simple_explanation': "Bagian-bagian tertentu dari frame ini menunjukkan 'jejak' editing digital. Seperti sidik jari pada kaca, ELA dapat melihat area yang telah dimodifikasi karena memiliki tingkat kompresi yang berbeda.",
                            'severity': 'high' if max_diff > 100 else 'medium',
                            'implications': "Area dengan nilai ELA tinggi menunjukkan: (1) Objek yang disisipkan, (2) Area yang di-retouch, (3) Teks atau watermark yang ditambahkan"
                        }
                        f.evidence_obj.explanations['ela'] = ela_explanation
                        result.detailed_anomaly_analysis['compression_anomalies'].append(ela_explanation)

    # Konversi reasons list ke string untuk konsistensi
    for f in frames:
        if isinstance(f.evidence_obj.reasons, list) and f.evidence_obj.reasons:
            f.evidence_obj.reasons = ", ".join(sorted(list(set(f.evidence_obj.reasons))))

    # Analisis statistik keseluruhan
    log(f"\n  {Icons.ANALYSIS} ANALISIS STATISTIK KESELURUHAN...")

    # Hitung distribusi anomali
    anomaly_types = Counter()
    confidence_levels = Counter()
    temporal_distribution = []

    for f in frames:
        if f.type.startswith("anomaly"):
            anomaly_types[f.type] += 1
            confidence_levels[f.evidence_obj.confidence] += 1
            temporal_distribution.append(f.timestamp)

    # Analisis clustering temporal anomali
    temporal_clusters = []
    if temporal_distribution:
        current_cluster = [temporal_distribution[0]]

        for i in range(1, len(temporal_distribution)):
            if temporal_distribution[i] - temporal_distribution[i-1] < 2.0:  # Within 2 seconds
                current_cluster.append(temporal_distribution[i])
            else:
                if len(current_cluster) > 1:
                    temporal_clusters.append(current_cluster)
                current_cluster = [temporal_distribution[i]]

        if len(current_cluster) > 1:
            temporal_clusters.append(current_cluster)

        log(f"  📊 Distribusi Anomali:")
        for atype, count in anomaly_types.items():
            log(f"     - {atype.replace('anomaly_', '').title()}: {count} frame")

        log(f"  📊 Tingkat Kepercayaan:")
        for level, count in confidence_levels.items():
            log(f"     - {level}: {count} anomali")

        if temporal_clusters:
            log(f"  📊 Ditemukan {len(temporal_clusters)} kluster anomali temporal")
            for i, cluster in enumerate(temporal_clusters):
                log(f"     - Kluster {i+1}: {len(cluster)} anomali dalam {cluster[-1]-cluster[0]:.2f} detik")

    # Simpan statistik dalam result
    result.statistical_summary = {
        'total_frames_analyzed': len(frames),
        'total_anomalies': sum(anomaly_types.values()),
        'anomaly_types': dict(anomaly_types),
        'confidence_distribution': dict(confidence_levels),
        'temporal_clusters': len(temporal_clusters) if temporal_distribution else 0,
        'average_anomalies_per_cluster': np.mean([len(c) for c in temporal_clusters]) if temporal_clusters else 0
    }

    # Update confidence distribution untuk Tahap 4
    result.confidence_distribution = dict(confidence_levels)

    # Buat visualisasi ringkasan anomali
    if anomaly_types:
        create_anomaly_summary_visualization(result, detail_viz_dir)

    log(f"\n  {Icons.SUCCESS} Tahap 3 Selesai - Investigasi mendalam telah dilengkapi dengan penjelasan detail.")

# Fungsi helper untuk membuat visualisasi ELA detail
def create_ela_visualization(original_path: Path, ela_array: np.ndarray, regional_analysis: dict, out_dir: Path) -> Path | None:
    """Membuat visualisasi ELA dengan highlight area mencurigakan."""
    try:
        # Load original image
        original = cv2.imread(str(original_path))
        if original is None:
            return None

        # Convert ELA array to color
        ela_color = cv2.applyColorMap((ela_array.mean(axis=2) * 5).astype(np.uint8), cv2.COLORMAP_JET)

        # Create combined visualization
        height, width = original.shape[:2]
        combined = np.zeros((height, width * 2 + 20, 3), dtype=np.uint8)
        combined[:, :width] = original
        combined[:, width+20:] = ela_color

        # Draw suspicious regions
        for region in regional_analysis['suspicious_regions'][:10]:  # Top 10
            x, y = region['x'], region['y']
            w, h = region['width'], region['height']
            color = (0, 0, 255) if region['suspicion_level'] == 'high' else (0, 255, 255)

            # Draw on original
            cv2.rectangle(combined, (x, y), (x+w, y+h), color, 2)
            # Draw on ELA
            cv2.rectangle(combined, (width+20+x, y), (width+20+x+w, y+h), color, 2)

        # Add labels
        font = cv2.FONT_HERSHEY_SIMPLEX
        cv2.putText(combined, 'Original', (10, 30), font, 1, (255, 255, 255), 2)
        cv2.putText(combined, 'ELA Analysis', (width+30, 30), font, 1, (255, 255, 255), 2)
        cv2.putText(combined, f'Suspicious Areas: {regional_analysis["suspicious_count"]}',
                    (10, height-10), font, 0.7, (255, 255, 0), 2)

        # Save
        out_path = out_dir / f"ela_detailed_{original_path.stem}.jpg"
        cv2.imwrite(str(out_path), combined)
        return out_path
    except Exception as e:
        log(f"  {Icons.ERROR} Error creating ELA visualization: {e}")
        return None

# Fungsi untuk membuat visualisasi ringkasan anomali
def create_anomaly_summary_visualization(result: AnalysisResult, out_dir: Path):
    """Membuat visualisasi ringkasan dari semua anomali yang terdeteksi."""
    try:
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 10))
        fig.suptitle('Ringkasan Analisis Forensik Video', fontsize=16, fontweight='bold')

        # 1. Pie chart distribusi tipe anomali
        if result.statistical_summary.get('anomaly_types'):
            labels = [t.replace('anomaly_', '').title() for t in result.statistical_summary['anomaly_types'].keys()]
            sizes = list(result.statistical_summary['anomaly_types'].values())
            colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99']

            ax1.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
            ax1.set_title('Distribusi Jenis Anomali')
            ax1.axis('equal')
        else:
            ax1.text(0.5, 0.5, 'Tidak ada anomali terdeteksi', ha='center', va='center')
            ax1.set_xlim(0, 1)
            ax1.set_ylim(0, 1)
            ax1.axis('off')

        # 2. Bar chart tingkat kepercayaan
        if result.statistical_summary.get('confidence_distribution'):
            confidence_labels = list(result.statistical_summary['confidence_distribution'].keys())
            confidence_values = list(result.statistical_summary['confidence_distribution'].values())
            colors_conf = {'RENDAH': 'green', 'SEDANG': 'yellow', 'TINGGI': 'orange', 'SANGAT TINGGI': 'red', 'N/A': 'gray'}
            bar_colors = [colors_conf.get(label, 'gray') for label in confidence_labels]

            ax2.bar(confidence_labels, confidence_values, color=bar_colors)
            ax2.set_title('Distribusi Tingkat Kepercayaan Anomali')
            ax2.set_xlabel('Tingkat Kepercayaan')
            ax2.set_ylabel('Jumlah Anomali')

        # 3. Timeline anomali
        anomaly_times = []
        anomaly_types_list = []
        for f in result.frames:
            if f.type.startswith("anomaly"):
                anomaly_times.append(f.timestamp)
                anomaly_types_list.append(f.type.replace('anomaly_', ''))

        if anomaly_times:
            # Create scatter plot with different colors for each type
            type_colors = {'discontinuity': 'purple', 'duplication': 'orange', 'insertion': 'red'}
            for atype in set(anomaly_types_list):
                times = [t for t, at in zip(anomaly_times, anomaly_types_list) if at == atype]
                ax3.scatter(times, [1]*len(times), label=atype.title(),
                           color=type_colors.get(atype, 'gray'), s=100, alpha=0.7)

            ax3.set_title('Timeline Anomali')
            ax3.set_xlabel('Waktu (detik)')
            ax3.set_ylim(0.5, 1.5)
            ax3.set_yticks([])
            ax3.legend()
            ax3.grid(True, axis='x', alpha=0.3)
        else:
            ax3.text(0.5, 0.5, 'Tidak ada timeline anomali', ha='center', va='center')
            ax3.set_xlim(0, 1)
            ax3.set_ylim(0, 1)
            ax3.axis('off')

        # 4. Statistik ringkasan
        stats_text = f"""Total Frame Dianalisis: {result.statistical_summary.get('total_frames_analyzed', 'N/A')}
Total Anomali Terdeteksi: {result.statistical_summary.get('total_anomalies', 'N/A')}
Persentase Anomali: {result.statistical_summary.get('total_anomalies', 0)/result.statistical_summary.get('total_frames_analyzed', 1)*100:.1f}%
Kluster Temporal: {result.statistical_summary.get('temporal_clusters', 'N/A')}
Rata-rata Anomali per Kluster: {result.statistical_summary.get('average_anomalies_per_cluster', 0):.1f}"""

        ax4.text(0.1, 0.5, stats_text, fontsize=12, verticalalignment='center',
                fontfamily='monospace', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
        ax4.set_xlim(0, 1)
        ax4.set_ylim(0, 1)
        ax4.axis('off')
        ax4.set_title('Statistik Ringkasan')

        plt.tight_layout(rect=[0, 0.03, 1, 0.95])
        summary_path = out_dir / "anomaly_summary.png"
        plt.savefig(summary_path, dpi=150, bbox_inches='tight')
        plt.close()

        result.plots['anomaly_summary'] = str(summary_path)
    except Exception as e:
        log(f"  {Icons.ERROR} Error creating summary visualization: {e}")

# --- TAHAP 4: VISUALISASI & PENILAIAN INTEGRITAS (ENHANCED VERSION) ---
def run_tahap_4_visualisasi_dan_penilaian(result: AnalysisResult, out_dir: Path):
    print_stage_banner(4, "Visualisasi & Penilaian Keandalan Bukti", "📊",
                       "Membuat plot detail, melokalisasi peristiwa, menilai keandalan bukti dengan FERM, dan menilai pipeline.")

    log(f"  {Icons.ANALYSIS} METODE UTAMA: Melakukan Localization Tampering untuk mengelompokkan anomali...")
    log(f"  📖 Localization Tampering adalah teknik untuk mengelompokkan frame-frame anomali yang berdekatan")
    log(f"     menjadi satu 'peristiwa' yang koheren, memudahkan interpretasi hasil forensik.")

    locs, event = [], None
    for f in result.frames:
        is_anomaly = f.type.startswith("anomaly")
        if is_anomaly:
            image_to_show = f.img_path_original
            if event and event["event"] == f.type and f.index == event["end_frame"] + 1:
                # Extend existing event
                event["end_frame"] = f.index
                event["end_ts"] = f.timestamp
                event["frame_count"] += 1
                # Update confidence ke yang tertinggi
                conf_hierarchy = {"SANGAT TINGGI": 4, "TINGGI": 3, "SEDANG": 2, "RENDAH": 1, "N/A": 0}
                if conf_hierarchy.get(f.evidence_obj.confidence, 0) > conf_hierarchy.get(event["confidence"], 0):
                    event["confidence"] = f.evidence_obj.confidence
                # Update explanations
                if f.evidence_obj.explanations:
                    event["explanations"].update(f.evidence_obj.explanations)
                # Collect all metrics
                event["all_metrics"].append(f.evidence_obj.metrics)
            else:
                # Save previous event if exists
                if event:
                    locs.append(event)
                # Start new event
                event = {
                    "event": f.type,
                    "start_frame": f.index,
                    "end_frame": f.index,
                    "start_ts": f.timestamp,
                    "end_ts": f.timestamp,
                    "frame_count": 1,
                    "confidence": f.evidence_obj.confidence,
                    "reasons": str(f.evidence_obj.reasons),
                    "metrics": f.evidence_obj.metrics,
                    "all_metrics": [f.evidence_obj.metrics],  # Collect all metrics for statistics
                    "image": image_to_show,
                    "ela_path": f.evidence_obj.ela_path,
                    "sift_path": f.evidence_obj.sift_path,
                    "explanations": f.evidence_obj.explanations.copy(),
                    "visualizations": f.evidence_obj.visualizations.copy()
                }
        elif event:
            locs.append(event)
            event = None
    if event:
        locs.append(event)

    # Enhance localization dengan analisis tambahan
    for loc in locs:
        # Calculate event duration and severity
        loc['duration'] = loc['end_ts'] - loc['start_ts']
        loc['severity_score'] = calculate_event_severity(loc)

        # Aggregate metrics across all frames in event
        if loc.get('all_metrics'):
            aggregated = {}
            for metrics in loc['all_metrics']:
                if isinstance(metrics, dict):
                    for key, val in metrics.items():
                        if not isinstance(key, list) and not isinstance(val, list):
                            if key not in aggregated:
                                aggregated[key] = []
                            if val is not None:
                                aggregated[key].append(val)

            # Calculate statistics for numeric metrics
            loc['aggregated_metrics'] = {}
            for key, vals in aggregated.items():
                numeric_vals = [v for v in vals if isinstance(v, (int, float))]
                if numeric_vals:
                    loc['aggregated_metrics'][key] = {
                        'mean': np.mean(numeric_vals),
                        'max': max(numeric_vals),
                        'min': min(numeric_vals),
                        'std': np.std(numeric_vals)
                    }

    result.localizations = locs
    result.localization_details = {
        'total_events': len(locs),
        'events_by_type': Counter(loc['event'] for loc in locs),
        'total_anomalous_frames': sum(loc.get('frame_count', 0) for loc in locs),
        'average_event_duration': np.mean([loc.get('duration',0) for loc in locs]) if locs else 0,
        'max_event_duration': max([loc.get('duration',0) for loc in locs]) if locs else 0,
        'high_severity_events': sum(1 for loc in locs if loc.get('severity_score',0) > 0.7)
    }

    log(f"  -> Ditemukan dan dilokalisasi {len(locs)} peristiwa anomali.")
    log(f"  -> Rata-rata durasi peristiwa: {result.localization_details['average_event_duration']:.2f} detik")
    log(f"  -> Peristiwa dengan severity tinggi: {result.localization_details['high_severity_events']}")

    # Calculate comprehensive summary
    total_anom = sum(1 for f in result.frames if f.type.startswith("anomaly"))
    total_frames = len(result.frames)
    pct_anomaly = round(total_anom * 100 / total_frames, 2) if total_frames > 0 else 0
    result.summary = {
        "total_frames": total_frames,
        "total_anomaly": total_anom,
        "pct_anomaly": pct_anomaly,
        "total_events": len(locs),
        "anomaly_density": total_anom / total_frames if total_frames > 0 else 0
    }

    log(f"  {Icons.INFO} {total_anom} dari {total_frames} frame terindikasi anomali ({pct_anomaly}%).")

    # Menghasilkan Forensic Evidence Reliability Matrix
    log(f"\n  {Icons.ANALYSIS} Menghasilkan Forensic Evidence Reliability Matrix (FERM)...")
    ferm_results = generate_forensic_evidence_matrix(result)
    result.forensic_evidence_matrix = ferm_results

    # Buat visualisasi FERM
    log(f"  📊 Membuat visualisasi matriks bukti forensik...")
    ferm_viz_paths = create_ferm_visualizations(result, ferm_results, out_dir)
    for viz_type, path in ferm_viz_paths.items():
        result.plots[f'ferm_{viz_type}'] = path

    log(f"  -> Penilaian Reliabilitas: {ferm_results['conclusion']['reliability_assessment']}")
    log(f"  -> Jumlah temuan utama: {len(ferm_results['conclusion']['primary_findings'])}")
    log(f"  -> Faktor reliabilitas: {len(ferm_results['conclusion']['reliability_factors'])}")

    # Assess pipeline performance
    log(f"\n  {Icons.EXAMINATION} Menilai performa setiap tahap pipeline forensik...")
    result.pipeline_assessment = assess_pipeline_performance(result)
    for stage_id, assessment in result.pipeline_assessment.items():
        log(f"  -> {assessment['nama']}: Quality Score = {assessment['quality_score']}%")

    # Create enhanced visualizations
    log(f"\n  {Icons.ANALYSIS} Membuat visualisasi detail...")

    # 1. Enhanced Localization Map
    log(f"  📍 Membuat peta lokalisasi tampering yang detail...")
    enhanced_map_path = create_enhanced_localization_map(result, out_dir)
    result.plots['enhanced_localization_map'] = str(enhanced_map_path)

    # 2. Anomaly Explanation Infographic
    log(f"  📚 Membuat infografis penjelasan anomali untuk orang awam...")
    infographic_path = create_anomaly_explanation_infographic(result, out_dir)
    result.plots['anomaly_infographic'] = str(infographic_path)

    # 3. Existing plots (dengan perbaikan)
    log(f"  📈 Membuat plot temporal standar...")

    # K-Means temporal plot
    color_clusters = [f.color_cluster for f in result.frames if f.color_cluster is not None]
    if color_clusters:
        plt.figure(figsize=(15, 6))
        plt.plot(range(len(color_clusters)), color_clusters, marker='.', linestyle='-', markersize=4, label='Klaster Warna Frame')
        jump_frames = [i for i in range(1, len(color_clusters)) if color_clusters[i] != color_clusters[i-1]]
        if jump_frames:
            for jf in jump_frames:
                plt.axvline(x=jf, color='r', linestyle='--', linewidth=1, alpha=0.7)
            plt.plot([], [], color='r', linestyle='--', linewidth=1, label='Perubahan Adegan Terdeteksi')
        plt.title('Visualisasi Klasterisasi Warna (Metode K-Means) Sepanjang Waktu', fontsize=14, weight='bold')
        plt.xlabel('Indeks Frame', fontsize=12)
        plt.ylabel('Nomor Klaster Warna', fontsize=12)
        if len(set(color_clusters)) > 1:
            plt.yticks(range(min(set(color_clusters)), max(set(color_clusters))+1))
        plt.grid(True, linestyle=':', alpha=0.7)
        plt.legend(loc='upper right', fontsize=10)
        plt.tight_layout()
        kmeans_temporal_plot_path = out_dir / f"plot_kmeans_temporal_{Path(result.video_path).stem}.png"
        plt.savefig(kmeans_temporal_plot_path, bbox_inches="tight", dpi=150)
        plt.close()
        result.plots['kmeans_temporal'] = str(kmeans_temporal_plot_path)

    # SSIM temporal plot
    ssim_values = [f.ssim_to_prev for f in result.frames if f.ssim_to_prev is not None]
    if len(ssim_values) > 1:
        y_values_ssim = ssim_values[1:]
        x_indices_ssim = list(range(1, len(y_values_ssim) + 1))

        plt.figure(figsize=(15, 6))
        plt.plot(x_indices_ssim, y_values_ssim, color='skyblue', marker='.', linestyle='-', markersize=3, alpha=0.7)

        discontinuity_frames_ssim_indices = [f.index for f in result.frames if "SSIM" in str(f.evidence_obj.reasons)]
        if discontinuity_frames_ssim_indices:
            valid_indices = [i for i in discontinuity_frames_ssim_indices if 0 < i < len(ssim_values)]
            if valid_indices:
                discontinuity_ssim_y_values = [ssim_values[i] for i in valid_indices]
                plt.scatter(valid_indices, discontinuity_ssim_y_values, color='red', marker='X', s=100, zorder=5, label='Diskontinuitas Terdeteksi (SSIM)')

        plt.title('Perubahan SSIM Antar Frame Sepanjang Waktu', fontsize=14, weight='bold')
        plt.xlabel('Indeks Frame', fontsize=12)
        plt.ylabel('Skor SSIM (0-1, Lebih Tinggi Lebih Mirip)', fontsize=12)
        plt.ylim(0, 1.05)
        plt.grid(True, linestyle=':', alpha=0.7)
        plt.legend(loc='lower left', fontsize=10)
        plt.tight_layout()
        ssim_temporal_plot_path = out_dir / f"plot_ssim_temporal_{Path(result.video_path).stem}.png"
        plt.savefig(ssim_temporal_plot_path, bbox_inches="tight", dpi=150)
        plt.close()
        result.plots['ssim_temporal'] = str(ssim_temporal_plot_path)

    # Optical flow temporal plot
    flow_values = [f.optical_flow_mag for f in result.frames if f.optical_flow_mag is not None]
    if len(flow_values) > 1:
        y_values_flow = flow_values[1:]
        x_indices_flow = list(range(1, len(y_values_flow) + 1))

        plt.figure(figsize=(15, 6))
        plt.plot(x_indices_flow, y_values_flow, color='salmon', marker='.', linestyle='-', markersize=3, alpha=0.7)

        discontinuity_frames_flow_indices = [f.index for f in result.frames if "Aliran Optik" in str(f.evidence_obj.reasons)]
        if discontinuity_frames_flow_indices:
            valid_indices_flow = [i for i in discontinuity_frames_flow_indices if 0 < i < len(flow_values)]
            if valid_indices_flow:
                discontinuity_flow_y_values = [flow_values[i] for i in valid_indices_flow]
                plt.scatter(valid_indices_flow, discontinuity_flow_y_values, color='darkgreen', marker='o', s=100, zorder=5, label='Diskontinuitas Terdeteksi (Aliran Optik)')

        flow_mags_for_z = [m for m in flow_values if m is not None and m > 0.0]
        if len(flow_mags_for_z) > 1:
            median_flow = np.median(flow_mags_for_z)
            mad_flow = stats.median_abs_deviation(flow_mags_for_z)
            mad_flow = 1e-9 if mad_flow == 0 else mad_flow
            threshold_mag_upper = (CONFIG["OPTICAL_FLOW_Z_THRESH"] / 0.6745) * mad_flow + median_flow
            plt.axhline(y=threshold_mag_upper, color='blue', linestyle='--', linewidth=1, label=f'Ambang Batas Atas Z-score')

        plt.title('Perubahan Rata-rata Magnitudo Aliran Optik', fontsize=14, weight='bold')
        plt.xlabel('Indeks Frame', fontsize=12)
        plt.ylabel('Rata-rata Magnitudo Aliran Optik', fontsize=12)
        plt.grid(True, linestyle=':', alpha=0.7)
        plt.legend(loc='upper right', fontsize=10)
        plt.tight_layout()
        optical_flow_temporal_plot_path = out_dir / f"plot_optical_flow_temporal_{Path(result.video_path).stem}.png"
        plt.savefig(optical_flow_temporal_plot_path, bbox_inches="tight", dpi=150)
        plt.close()
        result.plots['optical_flow_temporal'] = str(optical_flow_temporal_plot_path)

    # Metrics histograms
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4))
    if len(ssim_values) > 1:
        ssim_to_plot = [s for s in ssim_values[1:] if s is not None]
        if ssim_to_plot:
            ax1.hist(ssim_to_plot, bins=50, color='skyblue', edgecolor='black')
        ax1.set_title("Distribusi Skor SSIM")
        ax1.set_xlabel("Skor SSIM")
        ax1.set_ylabel("Frekuensi")
    if len(flow_values) > 1:
        flow_to_plot = [f for f in flow_values[1:] if f is not None]
        if flow_to_plot:
            ax2.hist(flow_to_plot, bins=50, color='salmon', edgecolor='black')
        ax2.set_title("Distribusi Aliran Optik")
        ax2.set_xlabel("Rata-rata Pergerakan")
        ax2.set_ylabel("Frekuensi")
    plt.tight_layout()
    metrics_histograms_plot_path = out_dir / f"plot_metrics_histograms_{Path(result.video_path).stem}.png"
    plt.savefig(metrics_histograms_plot_path, dpi=100)
    plt.close()
    result.plots['metrics_histograms'] = str(metrics_histograms_plot_path)

    # Simple temporal anomaly plot
    plt.figure(figsize=(15, 6))
    anomaly_data = {
        'Duplikasi': {'x': [], 'color': 'orange', 'marker': 'o', 'level': 1.0},
        'Penyisipan': {'x': [], 'color': 'red', 'marker': 'x', 'level': 0.9},
        'Diskontinuitas': {'x': [], 'color': 'purple', 'marker': '|', 'level': 0.8}
    }
    for f in result.frames:
        if f.type == "anomaly_duplication":
            anomaly_data['Duplikasi']['x'].append(f.index)
        elif f.type == "anomaly_insertion":
            anomaly_data['Penyisipan']['x'].append(f.index)
        elif f.type == "anomaly_discontinuity":
            anomaly_data['Diskontinuitas']['x'].append(f.index)

    for label, data in anomaly_data.items():
        if data['x']:
            plt.vlines(data['x'], 0, data['level'], colors=data['color'], lw=1.5, alpha=0.8)
            plt.scatter(data['x'], np.full_like(data['x'], data['level'], dtype=float),
                       c=data['color'], marker=data['marker'], s=40, label=label, zorder=5)

    plt.ylim(-0.1, 1.2)
    plt.yticks([0, 0.8, 0.9, 1.0], ['Asli', 'Diskontinuitas', 'Penyisipan', 'Duplikasi'])
    plt.xlabel("Indeks Frame", fontsize=12)
    plt.ylabel("Jenis Anomali Terdeteksi", fontsize=12)
    plt.title(f"Peta Anomali Temporal untuk {Path(result.video_path).name}", fontsize=14, weight='bold')
    plt.grid(True, axis='x', linestyle=':', alpha=0.7)

    from matplotlib.lines import Line2D
    plt.legend(handles=[Line2D([0], [0], color=d['color'], marker=d['marker'], linestyle='None', label=l)
                        for l, d in anomaly_data.items() if d['x']], loc='upper right', fontsize=10)
    plt.tight_layout()
    temporal_plot_path = out_dir / f"plot_temporal_{Path(result.video_path).stem}.png"
    plt.savefig(temporal_plot_path, bbox_inches="tight", dpi=150)
    plt.close()
    result.plots['temporal'] = str(temporal_plot_path)

    log(f"  {Icons.SUCCESS} Tahap 4 Selesai - Analisis detail dan penilaian integritas telah lengkap.")

# Helper function for calculating event severity
def calculate_event_severity(event: dict) -> float:
    """Calculate severity score for an anomaly event (0-1)."""
    severity = 0.0

    # Base severity by type
    type_severity = {
        'anomaly_insertion': 0.8,
        'anomaly_duplication': 0.6,
        'anomaly_discontinuity': 0.5
    }
    severity = type_severity.get(event.get('event', ''), 0.3)

    # Adjust by confidence
    confidence_multiplier = {
        'SANGAT TINGGI': 1.2,
        'TINGGI': 1.0,
        'SEDANG': 0.8,
        'RENDAH': 0.6,
        'N/A': 0.5
    }
    severity *= confidence_multiplier.get(event.get('confidence', 'N/A'), 0.5)

    # Adjust by duration (longer events are more severe)
    duration = event.get('duration', 0)
    if duration > 5.0:
        severity *= 1.2
    elif duration > 2.0:
        severity *= 1.1

    # Adjust by frame count
    frame_count = event.get('frame_count', 0)
    if frame_count > 10:
        severity *= 1.1

    # Normalize to 0-1 range
    severity = min(1.0, max(0.0, severity))

    return severity

# --- TAHAP 5: PENYUSUNAN LAPORAN & VALIDASI FORENSIK ---
def run_tahap_5_pelaporan_dan_validasi(result, out_dir, baseline_result, include_simple: bool = True, include_technical: bool = True):
    """Generate laporan PDF dengan layout konsisten menggunakan WeasyPrint dan template HTML"""

    # 1. Sanitasi input HTML/Markdown
    def sanitize_content(content):
        """Konversi Markdown → HTML dan bersihkan tag tidak aman"""
        html = markdown2.markdown(content, extras=["tables"])
        return bleach.clean(
            html,
            tags=['p', 'h1', 'h2', 'h3', 'table', 'tr', 'td', 'th', 'ul', 'ol', 'li', 'strong', 'em', 'img', 'div', 'span'],
            attributes={'*': ['style', 'class', 'src', 'alt']},
            protocols=['data', 'http', 'https'],
            strip=True
        )

    # 2. Persiapkan konten laporan
    report_data = {
        'judul': 'Laporan Analisis Forensik Video',
        'tanggal': datetime.now().strftime("%d %B %Y"),
        'preservation_hash': result.preservation_hash,
        'video_name': Path(result.video_path).name,
        'total_frames': result.summary.get('total_frames', 'N/A'),
        'total_anomalies': result.summary.get('total_anomaly', 'N/A'),
        'pct_anomaly': result.summary.get('pct_anomaly', 'N/A'),
        'total_events': result.localization_details.get('total_events', 'N/A'),
        'reliability_assessment': result.forensic_evidence_matrix['conclusion']['reliability_assessment'],
        'primary_findings': result.forensic_evidence_matrix['conclusion']['primary_findings'],
        'recommended_actions': result.forensic_evidence_matrix['conclusion']['recommended_actions'],
        'localizations': [],  # akan diisi dengan bytes
        'plots': {},          # akan diisi dengan bytes
        'include_simple': include_simple,
        'include_technical': include_technical,
        'metadata': result.metadata,
        'metadata_analysis': result.metadata_analysis
    }

    # Konversi gambar lokal dan plot menjadi bytes untuk penyematan
    locs_for_report: list[dict] = []
    for loc in result.localizations:
        loc_copy = loc.copy()
        for fkey in ("image", "ela_path", "sift_path"):
            fpath = loc_copy.get(fkey)
            try:
                if fpath and Path(fpath).exists():
                    loc_copy[f"{fkey}_bytes"] = Path(fpath).read_bytes()
            except Exception as e:
                log(f"  {Icons.ERROR} Gagal membaca {fkey} {fpath}: {e}")
        locs_for_report.append(loc_copy)
    report_data['localizations'] = locs_for_report

    plot_bytes = {}
    for key, pth in result.plots.items():
        try:
            if pth and Path(pth).exists():
                plot_bytes[f"{key}_bytes"] = Path(pth).read_bytes()
        except Exception as e:
            log(f"  {Icons.ERROR} Gagal membaca plot {key}: {e}")
    report_data['plots'] = plot_bytes

    # 3. Template HTML dengan CSS styling
    template_str = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{{ judul }}</title>
        <style>
            body { font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.5; color: #333; }
            h1, h2, h3, h4 { margin-top: 1.2em; margin-bottom: 0.5em; text-align: center; color: #0B3D91; }
            h1 { font-size: 24pt; }
            h2 { font-size: 18pt; }
            h3 { font-size: 14pt; }
            table { width: 100%; border-collapse: collapse; margin: 1em 0; }
            th, td { border: 0.5px solid #ccc; padding: 6px; vertical-align: top; text-align: left; }
            th { background-color: #f2f2f2; }
            img { max-width: 90%; height: auto; display: block; margin: 1em auto; border: 1px solid #eee; box-shadow: 2px 2px 5px rgba(0,0,0,0.1); }
            .footer { text-align: center; font-size: 8pt; margin-top: 1em; }
            .section { margin-bottom: 2em; padding: 1em; border: 1px solid #eee; border-radius: 5px; background-color: #fff; }
            .reliability-high { color: #155724; background-color: #d4edda; border-color: #c3e6cb; padding: 5px; border-radius: 3px; font-weight: bold; text-align: center; }
            .reliability-moderate { color: #856404; background-color: #fff3cd; border-color: #ffeeba; padding: 5px; border-radius: 3px; font-weight: bold; text-align: center; }
            .reliability-limited { color: #721c24; background-color: #f8d7da; border-color: #f5c6cb; padding: 5px; border-radius: 3px; font-weight: bold; text-align: center; }
            .reliability-low { color: #721c24; background-color: #f8d7da; border-color: #f5c6cb; padding: 5px; border-radius: 3px; font-weight: bold; text-align: center; }
            .finding-box { background-color: #f0f7ff; border-left: 4px solid #0c6dd6; padding: 10px; margin-bottom: 10px; border-radius: 3px; }
            .anomaly-event-box { background-color: #f8f9fa; border-left: 5px solid #0c6dd6; padding: 15px; margin-bottom: 15px; border-radius: 5px; }
            .anomaly-event-box h4 { text-align: left; color: #0056b3; }
            .metric-table { width: auto; margin: 0.5em 0; font-size: 10pt; }
            .metric-table th, .metric-table td { padding: 3px 8px; }
            .explanation-simple { background-color: #FFF8DC; border: 1px solid #FFD700; padding: 8px; margin-top: 5px; border-radius: 3px; font-style: italic; }
            .explanation-technical { background-color: #F0F8FF; border: 1px solid #4682B4; padding: 8px; margin-top: 5px; border-radius: 3px; font-family: monospace; font-size: 10pt; }
        </style>
    </head>
    <body>
        <h1>{{ judul }}</h1>
        <p style="text-align:center">Tanggal Analisis: {{ tanggal }}</p>
        <p style="text-align:center">Nama Video: {{ video_name }}</p>
        <p style="text-align:center">Hash Preservasi (SHA-256): <code>{{ preservation_hash }}</code></p>

        <div class="section">
            <h2>Ringkasan Eksekutif</h2>
            <p>Laporan ini menyajikan hasil analisis forensik video menggunakan sistem VIFA-Pro. Tujuan analisis adalah untuk mengidentifikasi potensi manipulasi atau anomali dalam video bukti.</p>
            <h3>Statistik Utama</h3>
            <table>
                <tr><th>Metrik</th><th>Nilai</th></tr>
                <tr><td>Total Frame Dianalisis</td><td>{{ total_frames }}</td></tr>
                <tr><td>Total Anomali Terdeteksi</td><td>{{ total_anomalies }}</td></tr>
                <tr><td>Persentase Anomali</td><td>{{ pct_anomaly }}%</td></tr>
                <tr><td>Jumlah Peristiwa Anomali</td><td>{{ total_events }}</td></tr>
            </table>
            <h3>Penilaian Reliabilitas Bukti Forensik (FERM)</h3>
            <div class="reliability-{{ reliability_assessment.split(' ')[1]|lower }}">
                {{ reliability_assessment }}
            </div>
            <h4>Temuan Utama</h4>
            {% if primary_findings %}
                {% for finding in primary_findings %}
                <div class="finding-box">
                    <strong>{{ finding.finding }}</strong> (Kepercayaan: {{ finding.confidence }})<br>
                    <em>Bukti:</em> {{ finding.evidence }}<br>
                    <em>Interpretasi:</em> {{ finding.interpretation }}
                </div>
                {% endfor %}
            {% else %}
                <p>Tidak ada temuan utama yang teridentifikasi.</p>
            {% endif %}
            <h4>Rekomendasi Tindakan</h4>
            {% if recommended_actions %}
                <ul>
                {% for action in recommended_actions %}
                    <li>{{ action }}</li>
                {% endfor %}
                </ul>
            {% else %}
                <p>Tidak ada rekomendasi tindakan.</p>
            {% endif %}
        </div>

        <div class="section">
            <h2>Detail Peristiwa Anomali</h2>
            {% if localizations %}
                {% for loc in localizations %}
                <div class="anomaly-event-box">
                    <h4>Peristiwa: {{ loc.event|replace('anomaly_', '')|capitalize }} (Frame {{ loc.start_frame }} - {{ loc.end_frame }})</h4>
                    <p><strong>Durasi:</strong> {{ "%.2f"|format(loc.duration) }} detik</p>
                    <p><strong>Kepercayaan:</strong> {{ loc.confidence }}</p>
                    <p><strong>Alasan Deteksi:</strong> {{ loc.reasons }}</p>
                    
                    {% if include_simple and loc.explanations %}
                        {% for key, exp in loc.explanations.items() %}
                            {% if exp.simple_explanation %}
                                <div class="explanation-simple">
                                    <strong>Apa Artinya Ini?</strong> {{ exp.simple_explanation }}
                                </div>
                            {% endif %}
                        {% endfor %}
                    {% endif %}

                    {% if include_technical and loc.explanations %}
                        {% for key, exp in loc.explanations.items() %}
                            {% if exp.technical_explanation %}
                                <div class="explanation-technical">
                                    <strong>Detail Teknis:</strong> {{ exp.technical_explanation }}
                                    {% if exp.metrics %}
                                        <table class="metric-table">
                                            <tr><th colspan="2">Metrik Kunci</th></tr>
                                            {% for m_key, m_val in exp.metrics.items() %}
                                                <tr><td>{{ m_key|replace('_', ' ')|capitalize }}</td><td>{{ "%.3f"|format(m_val) if m_val is number else m_val }}</td></tr>
                                            {% endfor %}
                                        </table>
                                    {% endif %}
                                </div>
                            {% endif %}
                        {% endfor %}
                    {% endif %}

                    {% if loc.image %}
                        <p><strong>Frame Bukti Visual:</strong></p>
                        <img src="data:image/jpeg;base64,{{ loc.image|b64encode }}" alt="Frame Bukti">
                    {% endif %}
                    {% if loc.ela_path %}
                        <p><strong>Analisis ELA:</strong></p>
                        <img src="data:image/jpeg;base64,{{ loc.ela_path|b64encode }}" alt="ELA">
                    {% endif %}
                    {% if loc.sift_path %}
                        <p><strong>Analisis SIFT:</strong></p>
                        <img src="data:image/jpeg;base64,{{ loc.sift_path|b64encode }}" alt="SIFT">
                    {% endif %}
                </div>
                {% endfor %}
            {% else %}
                <p>Tidak ada peristiwa anomali signifikan yang terdeteksi.</p>
            {% endif %}
        </div>

        <div class="section">
            <h2>Visualisasi Hasil Analisis</h2>
            {% if plots %}
                {% for key, path in plots.items() %}
                    {% if key.endswith('_bytes') %}
                        <h3>{{ key|replace('_bytes', '')|replace('_', ' ')|capitalize }}</h3>
                        <img src="data:image/png;base64,{{ path|b64encode }}" alt="{{ key }}">
                    {% endif %}
                {% endfor %}
            {% else %}
                <p>Tidak ada visualisasi yang dihasilkan.</p>
            {% endif %}
        </div>

        <div class="section">
            <h2>Metadata Video</h2>
            {% if metadata %}
                {% for category, items in metadata.items() %}
                    <h3>{{ category }}</h3>
                    <table>
                        {% for key, value in items.items() %}
                            <tr><td>{{ key }}</td><td>{{ value }}</td></tr>
                        {% endfor %}
                    </table>
                {% endfor %}
            {% else %}
                <p>Tidak ada metadata yang tersedia.</p>
            {% endif %}
        </div>

        <div class="section">
            <h2>Analisis Metadata</h2>
            {% if metadata_analysis %}
                <table>
                    <tr><td>Status Editing</td><td>{{ 'Ya' if metadata_analysis.edited else 'Tidak' }}</td></tr>
                    <tr><td>Perangkat Lunak</td><td>{{ metadata_analysis.editing_software }}</td></tr>
                    {% if metadata_analysis.creation_time %}
                        <tr><td>Waktu Pembuatan</td><td>{{ metadata_analysis.creation_time }}</td></tr>
                    {% endif %}
                </table>
            {% else %}
                <p>Tidak ada analisis metadata.</p>
            {% endif %}
        </div>

        <div class="footer">
            <p>Laporan ini dihasilkan secara otomatis oleh VIFA-Pro.</p>
        </div>
    </body>
    </html>
    """

    # 4. Render template dengan data
    env = Environment(loader=BaseLoader())
    # Add base64 filter yang dapat menerima bytes atau path file
    import base64
    def b64encode_data(data):
        try:
            if isinstance(data, bytes):
                raw = data
            else:
                path_obj = Path(str(data))
                raw = path_obj.read_bytes()
            return base64.b64encode(raw).decode('utf-8')
        except Exception:
            return ''

    env.filters['b64encode'] = b64encode_data
    template = env.from_string(template_str)
    html_content = template.render(**report_data)

    # 5. Generate PDF dengan WeasyPrint
    html = HTML(string=html_content)
    css = CSS(string=f'''
        @page {{
            size: A4;
            margin: 2cm;
            @top-center {{
                content: "{report_data['judul']}";
                font-size: 10pt;
                color: #555;
            }}
            @bottom-center {{
                content: "Halaman " counter(page);
                font-size: 8pt;
                color: #777;
            }}
        }}
    ''')

    # 6. Simpan PDF ke buffer dan file
    pdf_buffer = io.BytesIO()
    html.write_pdf(pdf_buffer, stylesheets=[css])
    result.pdf_report_bytes = pdf_buffer.getvalue()

    report_path = Path(out_dir) / "laporan_forensik.pdf"
    with open(report_path, "wb") as f:
        f.write(result.pdf_report_bytes)
    result.pdf_report_path = str(report_path)

    # Generate DOCX report if available
    if DOCX_AVAILABLE:
        log(f"  {Icons.INFO} Menyusun laporan DOCX...")
        docx_path = Path(out_dir) / "laporan_forensik.docx"
        try:
            document = Document()
            styles = document.styles

            # Add custom styles
            if 'Normal' not in styles:
                styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)

            # Title
            document.add_heading(report_data['judul'], level=0)
            document.add_paragraph(f"Tanggal Analisis: {report_data['tanggal']}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(f"Nama Video: {report_data['video_name']}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(f"Hash Preservasi (SHA-256): {report_data['preservation_hash']}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_page_break()

            # Executive Summary
            document.add_heading('Ringkasan Eksekutif', level=1)
            document.add_paragraph('Laporan ini menyajikan hasil analisis forensik video menggunakan sistem VIFA-Pro. Tujuan analisis adalah untuk mengidentifikasi potensi manipulasi atau anomali dalam video bukti.')

            document.add_heading('Statistik Utama', level=2)
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = 'Metrik'
            table.rows[0].cells[1].text = 'Nilai'
            table.add_row().cells[0].text = 'Total Frame Dianalisis'; table.rows[1].cells[1].text = str(report_data['total_frames'])
            table.add_row().cells[0].text = 'Total Anomali Terdeteksi'; table.rows[2].cells[1].text = str(report_data['total_anomalies'])
            table.add_row().cells[0].text = 'Persentase Anomali'; table.rows[3].cells[1].text = f"{report_data['pct_anomaly']}%"
            table.add_row().cells[0].text = 'Jumlah Peristiwa Anomali'; table.rows[4].cells[1].text = str(report_data['total_events'])

            document.add_heading('Penilaian Reliabilitas Bukti Forensik (FERM)', level=2)
            p = document.add_paragraph()
            run = p.add_run(report_data['reliability_assessment'])
            if "Tinggi" in report_data['reliability_assessment']:
                run.font.color.rgb = RGBColor(0x15, 0x57, 0x24)
            elif "Sedang" in report_data['reliability_assessment']:
                run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
            else:
                run.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)
            run.bold = True

            document.add_heading('Temuan Utama', level=3)
            if report_data['primary_findings']:
                for finding in report_data['primary_findings']:
                    p = document.add_paragraph()
                    p.add_run(f"{finding['finding']} (Kepercayaan: {finding['confidence']})").bold = True
                    p.add_run(f"\nBukti: {finding['evidence']}")
                    p.add_run(f"\nInterpretasi: {finding['interpretation']}").italic = True
            else:
                document.add_paragraph('Tidak ada temuan utama yang teridentifikasi.')

            document.add_heading('Rekomendasi Tindakan', level=3)
            if report_data['recommended_actions']:
                for action in report_data['recommended_actions']:
                    document.add_paragraph(action, style='List Bullet')
            else:
                document.add_paragraph('Tidak ada rekomendasi tindakan.')
            document.add_page_break()

            # Anomaly Event Details
            document.add_heading('Detail Peristiwa Anomali', level=1)
            if report_data['localizations']:
                for loc in report_data['localizations']:
                    document.add_heading(f"Peristiwa: {loc['event'].replace('anomaly_', '').capitalize()} (Frame {loc['start_frame']} - {loc['end_frame']})", level=2)
                    document.add_paragraph(f"Durasi: {loc['duration']:.2f} detik")
                    document.add_paragraph(f"Kepercayaan: {loc['confidence']}")
                    document.add_paragraph(f"Alasan Deteksi: {loc['reasons']}")

                    if include_simple and loc.get('explanations'):
                        for key, exp in loc['explanations'].items():
                            if exp.get('simple_explanation'):
                                document.add_paragraph(f"Apa Artinya Ini? {exp['simple_explanation']}", style='Normal').italic = True

                    if include_technical and loc.get('explanations'):
                        for key, exp in loc['explanations'].items():
                            if exp.get('technical_explanation'):
                                document.add_paragraph(f"Detail Teknis: {exp['technical_explanation']}")
                                if exp.get('metrics'):
                                    metric_table = document.add_table(rows=1, cols=2)
                                    metric_table.rows[0].cells[0].text = 'Metrik Kunci'
                                    metric_table.rows[0].cells[1].text = 'Nilai'
                                    for m_key, m_val in exp['metrics'].items():
                                        row_cells = metric_table.add_row().cells
                                        row_cells[0].text = m_key.replace('_', ' ').capitalize()
                                        row_cells[1].text = f"{m_val:.3f}" if isinstance(m_val, (int, float)) else str(m_val)

                    if loc.get('image_bytes'):
                        document.add_paragraph('Frame Bukti Visual:')
                        image_stream = io.BytesIO(loc['image_bytes'])
                        document.add_picture(image_stream, width=Inches(5))
                    elif loc.get('image') and Path(loc['image']).exists():
                        document.add_paragraph('Frame Bukti Visual:')
                        document.add_picture(loc['image'], width=Inches(5))

                    if loc.get('ela_path_bytes'):
                        document.add_paragraph('Analisis ELA:')
                        image_stream = io.BytesIO(loc['ela_path_bytes'])
                        document.add_picture(image_stream, width=Inches(5))
                    elif loc.get('ela_path') and Path(loc['ela_path']).exists():
                        document.add_paragraph('Analisis ELA:')
                        document.add_picture(loc['ela_path'], width=Inches(5))

                    if loc.get('sift_path_bytes'):
                        document.add_paragraph('Analisis SIFT:')
                        image_stream = io.BytesIO(loc['sift_path_bytes'])
                        document.add_picture(image_stream, width=Inches(5))
                    elif loc.get('sift_path') and Path(loc['sift_path']).exists():
                        document.add_paragraph('Analisis SIFT:')
                        document.add_picture(loc['sift_path'], width=Inches(5))
                    document.add_paragraph() # Add a blank line for spacing
            else:
                document.add_paragraph('Tidak ada peristiwa anomali signifikan yang terdeteksi.')
            document.add_page_break()

            # Visualizations
            document.add_heading('Visualisasi Hasil Analisis', level=1)
            if report_data['plots']:
                for key, path_bytes in report_data['plots'].items():
                    if key.endswith('_bytes'):
                        document.add_heading(key.replace('_bytes', '').replace('_', ' ').capitalize(), level=2)
                        image_stream = io.BytesIO(path_bytes)
                        document.add_picture(image_stream, width=Inches(6))
                    else:
                        if Path(path_bytes).exists():
                            document.add_heading(key.replace('_', ' ').capitalize(), level=2)
                            document.add_picture(path_bytes, width=Inches(6))
            else:
                document.add_paragraph('Tidak ada visualisasi yang dihasilkan.')
            document.add_page_break()

            # Metadata
            document.add_heading('Metadata Video', level=1)
            if report_data['metadata']:
                for category, items in report_data['metadata'].items():
                    document.add_heading(category, level=2)
                    table = document.add_table(rows=1, cols=2)
                    table.rows[0].cells[0].text = 'Properti'
                    table.rows[0].cells[1].text = 'Nilai'
                    for key, value in items.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = key
                        row_cells[1].text = str(value)
            else:
                document.add_paragraph('Tidak ada metadata yang tersedia.')

            document.add_heading('Analisis Metadata', level=1)
            if report_data['metadata_analysis']:
                table = document.add_table(rows=1, cols=2)
                table.rows[0].cells[0].text = 'Kriteria'
                table.rows[0].cells[1].text = 'Nilai'
                table.add_row().cells[0].text = 'Status Editing'; table.rows[1].cells[1].text = 'Ya' if report_data['metadata_analysis']['edited'] else 'Tidak'
                table.add_row().cells[0].text = 'Perangkat Lunak'; table.rows[2].cells[1].text = report_data['metadata_analysis']['editing_software']
                if report_data['metadata_analysis'].get('creation_time'):
                    row = table.add_row().cells
                    row[0].text = 'Waktu Pembuatan'
                    row[1].text = str(report_data['metadata_analysis']['creation_time'])
            else:
                document.add_paragraph('Tidak ada analisis metadata.')

            document.save(docx_path)
            log(f"  ✅ Laporan DOCX berhasil dibuat: {docx_path.name}")
            result.docx_report_path = str(docx_path)
        except Exception as e:
            log(f"  {Icons.ERROR} Gagal membuat laporan DOCX: {e}")
            result.docx_report_path = None

    # 8. Buat arsip ZIP yang berisi seluruh laporan dan artefak
    zip_path = create_zip_archive(result, out_dir)
    if zip_path:
        result.zip_report_path = zip_path
        result.zip_report_bytes = zip_path.read_bytes()
    else:
        result.zip_report_path = None
        result.zip_report_bytes = None

    return result

###############################################################################
# MAIN EXECUTION
###############################################################################

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="VIFA-Pro: Sistem Forensik Video Profesional")
    parser.add_argument("video_path", type=str, help="Path ke video yang akan dianalisis")
    parser.add_argument("-b", "--baseline", type=str, help="Path ke video baseline (opsional)")
    parser.add_argument("-f", "--fps", type=int, default=10, help="FPS ekstraksi frame (default: 10)")
    parser.add_argument("-o", "--output", type=str, help="Direktori output (default: auto-generated)")

    args = parser.parse_args()

    video_path = Path(args.video_path)
    if not video_path.exists():
        print(f"{Icons.ERROR} File video tidak ditemukan: {video_path}")
        sys.exit(1)

    # Setup output directory
    if args.output:
        out_dir = Path(args.output)
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = Path(f"forensik_output_{video_path.stem}_{timestamp}")

    out_dir.mkdir(parents=True, exist_ok=True)

    # Run analysis
    print(f"\n{Icons.IDENTIFICATION} VIFA-Pro: Sistem Forensik Video Profesional")
    print(f"{Icons.INFO} Video: {video_path}")
    print(f"{Icons.INFO} Output: {out_dir}")
    print(f"{Icons.INFO} FPS: {args.fps}")

    # Main analysis
    result = run_tahap_1_pra_pemrosesan(video_path, out_dir, args.fps)
    if not result:
        print(f"{Icons.ERROR} Analisis gagal pada Tahap 1")
        sys.exit(1)

    # Baseline analysis if provided
    baseline_result = None
    if args.baseline:
        baseline_path = Path(args.baseline)
        if baseline_path.exists():
            print(f"\n{Icons.ANALYSIS} Memproses video baseline...")
            baseline_result = run_tahap_1_pra_pemrosesan(baseline_path, out_dir, args.fps)
            if baseline_result:
                run_tahap_2_analisis_temporal(baseline_result)

    # Continue with main analysis
    run_tahap_2_analisis_temporal(result, baseline_result)
    run_tahap_3_sintesis_bukti(result, out_dir)
    run_tahap_4_visualisasi_dan_penilaian(result, out_dir)
    run_tahap_5_pelaporan_dan_validasi(result, out_dir, baseline_result)

    print(f"\n{Icons.INFO} Penilaian Reliabilitas: {result.forensic_evidence_matrix['conclusion']['reliability_assessment']}")
    print(f"{Icons.INFO} Temuan Utama: {len(result.forensic_evidence_matrix['conclusion']['primary_findings'])}")

    print(f"\n{Icons.SUCCESS} Analisis selesai!")
    print(f"{Icons.INFO} Hasil tersimpan di: {out_dir}")
    print(f"{Icons.INFO} Laporan PDF: {result.pdf_report_path.name if result.pdf_report_path else 'N/A'}")

# --- END OF FILE ForensikVideo.py ---
