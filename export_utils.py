# --- START OF FILE export_utils.py ---

# export_utils.py
# VERSI BARU - FOKUS PADA STABILITAS PDF DAN DOCX
# REVISI UTAMA: Implementasi laporan DOCX yang komprehensif untuk menandingi detail PDF.

import io
from pathlib import Path
from typing import Optional, Any
import streamlit as st
from datetime import datetime
import shutil

# --- HELPER FUNCTIONS (direplikasi dari ForensikVideo untuk menghindari circular import) ---
def get_anomaly_explanation(event_type: str) -> str:
    explanations = {
        "Duplication": "Frame-frame ini adalah salinan identik dari frame sebelumnya. Dalam video asli, konten ini kemungkinan tidak diulang dan mungkin mengindikasikan manipulasi untuk memperpanjang durasi video atau menyembunyikan konten tertentu.",
        "Insertion": "Frame-frame ini <b>tidak ditemukan</b> dalam video asli/baseline. Ini mengindikasikan penambahan konten baru yang tidak ada pada rekaman original, yang mungkin bertujuan mengubah narasi atau konteks video.",
        "Discontinuity": "Terdeteksi 'patahan' atau transisi mendadak dalam aliran video. Hal ini mengindikasikan pemotongan bagian dari video asli, atau penyambungan konten dari sumber berbeda secara tidak mulus."
    }
    return explanations.get(event_type, "Jenis anomali tidak dikenal.")

def get_anomaly_implication(event_type: str) -> str:
    implications = {
        "Duplication": "Implikasi forensik dari duplikasi frame adalah kemungkinan adanya upaya untuk: (1) Memperpanjang durasi video secara artifisial, (2) Menutupi konten yang telah dihapus dengan mengulang konten yang ada, atau (3) Memanipulasi persepsi waktu dalam video tersebut.",
        "Insertion": "Penyisipan frame asing ke dalam video memiliki implikasi serius, termasuk: (1) Mengubah narasi atau konteks asli video, (2) Menambahkan elemen visual yang tidak ada pada saat perekaman asli, atau (3) Memalsukan bukti visual dengan menambahkan konten dari sumber lain.",
        "Discontinuity": "Diskontinuitas dalam video mengindikasikan: (1) Bagian tertentu dari video asli telah dihapus, (2) Konten dari sumber berbeda telah disambung secara tidak mulus, atau (3) Terjadi gangguan teknis selama proses pengambilan atau pengeditan video."
    }
    return implications.get(event_type, "Implikasi tidak dapat ditentukan untuk jenis anomali ini.")

def explain_metric(metric_name: str) -> str:
    explanations = {
        "optical_flow_z_score": "Ukuran lonjakan gerakan abnormal (Z-score > 4 = sangat abnormal).",
        "ssim_drop": "Ukuran penurunan kemiripan visual (> 0.25 = perubahan drastis).",
        "ssim_absolute_low": "Skor kemiripan yang sangat rendah (< 0.7 = sangat berbeda).",
        "color_cluster_jump": "Perubahan adegan visual berdasarkan analisis warna K-Means.",
        "source_frame": "Frame asli dari duplikasi (nomor indeks frame).",
        "ssim_to_source": "Skor kemiripan dengan frame asli (0-1, 1 = identik).",
        "sift_inliers": "Jumlah titik fitur unik yang cocok kuat (> 10 = duplikasi kuat).",
        "sift_good_matches": "Total kandidat titik fitur yang cocok.",
        "sift_inlier_ratio": "Rasio kecocokan valid (> 0.8 = duplikasi hampir pasti).",
        "ela_max_difference": "Tingkat perbedaan kompresi (0-255, > 100 = editing signifikan).",
        "ela_suspicious_regions": "Jumlah area yang menunjukkan tanda-tanda editing."
    }
    return explanations.get(metric_name, "Metrik analisis karakteristik visual/struktural frame.")

def check_dependency(package_name: str) -> bool:
    """Memeriksa apakah sebuah library Python terpasang."""
    import importlib.util
    return importlib.util.find_spec(package_name) is not None

def check_poppler_installation() -> bool:
    """Memeriksa ketersediaan utilitas Poppler `pdftoppm`."""
    return shutil.which("pdftoppm") is not None

def create_docx_report_robust(result: Any, output_path: Path) -> Optional[Path]:
    """
    Membuat laporan DOCX KOMPREHENSIF yang meniru detail laporan PDF.
    Struktur dan konten diperluas secara signifikan.
    """
    if not check_dependency('docx'):
        print("Peringatan: `python-docx` tidak terpasang. Ekspor DOCX dilewati.")
        return None

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        style = doc.styles['Normal'].font
        style.name = 'Arial'
        style.size = Pt(11)

        # Helper untuk shading box
        def shade_cell(cell, hex_color_string):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color_string)
            cell._tc.get_or_add_tcPr().append(shd)

        # --- Halaman Sampul ---
        doc.add_heading('Laporan Analisis Forensik Video', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('Dihasilkan oleh Sistem VIFA-Pro').italic = True
        
        doc.add_paragraph(f"Dihasilkan pada: {datetime.now().strftime('%d %B %Y, %H:%M:%S')}", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # --- Ringkasan Eksekutif & Metadata ---
        doc.add_heading('Ringkasan Eksekutif', level=1)
        
        # Info Metadata
        info_data = {
            "Nama File Bukti": Path(getattr(result, 'video_path', 'N/A')).name,
            "Hash Integritas (SHA-256)": getattr(result, 'preservation_hash', 'N/A'),
            "Total Frame Dianalisis": str(result.summary.get('total_frames', 'N/A')),
            "Total Anomali Ditemukan": str(result.summary.get('total_anomaly', 'N/A')),
        }
        table = doc.add_table(rows=1, cols=2, style='Table Grid')
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Detail"
        for key, value in info_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        doc.add_paragraph()
        
        # Penilaian FERM
        ferm = getattr(result, 'forensic_evidence_matrix', {})
        conclusion = ferm.get('conclusion', {})
        reliability = conclusion.get('reliability_assessment', 'Tidak Dapat Ditentukan')
        
        p = doc.add_paragraph()
        p.add_run('Penilaian Keandalan Keseluruhan: ').bold = True
        p.add_run(reliability)
        
        # Temuan Kunci dari FERM
        findings = conclusion.get('primary_findings', [])
        if findings:
            doc.add_paragraph('Temuan Kunci:', style='Intense Quote')
            for item in findings:
                finding_text = item.get('finding', 'N/A')
                confidence = item.get('confidence', 'N/A')
                doc.add_paragraph(f"• {finding_text} (Kepercayaan: {confidence})", style='List Bullet')
        
        # Rekomendasi
        recommendations = conclusion.get('recommended_actions', [])
        if recommendations:
            doc.add_paragraph('Rekomendasi Tindak Lanjut:', style='Intense Quote')
            for action in recommendations:
                doc.add_paragraph(f"• {action}", style='List Bullet')
        
        doc.add_page_break()

        # --- Detail Setiap Peristiwa Anomali (Meniru PDF) ---
        doc.add_heading('Detail Investigasi Setiap Peristiwa Anomali', level=1)
        localizations = getattr(result, 'localizations', [])
        if not localizations:
            doc.add_paragraph("Tidak ada peristiwa anomali signifikan yang ditemukan.")
        else:
            for i, loc in enumerate(localizations):
                event_type_raw = loc.get('event', 'unknown')
                event_type = event_type_raw.replace('anomaly_', '').capitalize()
                start_ts = loc.get('start_ts', 0)
                end_ts = loc.get('end_ts', 0)
                confidence = loc.get('confidence', 'N/A')
                
                doc.add_heading(f"Peristiwa #{i+1}: {event_type} @ {start_ts:.2f}s", level=2)
                p = doc.add_paragraph()
                p.add_run(f"Durasi: ").bold = True
                p.add_run(f"{end_ts - start_ts:.2f} detik | ")
                p.add_run(f"Kepercayaan: ").bold = True
                p.add_run(confidence)
                
                # Penjelasan Umum dan Implikasi
                doc.add_paragraph('Penjelasan Umum:', style='Intense Quote')
                doc.add_paragraph(get_anomaly_explanation(event_type))
                doc.add_paragraph('Implikasi Forensik:', style='Intense Quote')
                doc.add_paragraph(get_anomaly_implication(event_type))
                
                # Bukti Teknis
                doc.add_heading('Bukti Teknis Pendukung', level=3)
                if isinstance(loc.get('metrics'), dict):
                    metrics_table = doc.add_table(rows=1, cols=3, style='Table Grid')
                    metrics_table.cell(0,0).text = "Metrik"
                    metrics_table.cell(0,1).text = "Nilai"
                    metrics_table.cell(0,2).text = "Interpretasi"
                    for key, val in loc['metrics'].items():
                        row = metrics_table.add_row().cells
                        row[0].text = key.replace('_', ' ').title()
                        row[1].text = str(val)
                        row[2].text = explain_metric(key)
                else:
                    doc.add_paragraph("Tidak ada metrik teknis yang tercatat.")
                doc.add_paragraph()
                
                # Bukti Visual
                doc.add_heading('Bukti Visual', level=3)
                
                # Sample Frame
                if loc.get('image') and Path(loc['image']).exists():
                    doc.add_picture(loc['image'], width=Inches(3.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = doc.add_paragraph("Contoh Frame (Asli)", style='Caption')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # ELA
                if loc.get('ela_path') and Path(loc['ela_path']).exists():
                    doc.add_picture(loc['ela_path'], width=Inches(3.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = doc.add_paragraph("Analisis Kompresi (ELA)", style='Caption')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # SIFT
                if loc.get('sift_path') and Path(loc['sift_path']).exists():
                    doc.add_picture(loc['sift_path'], width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = doc.add_paragraph("Bukti Pencocokan Fitur SIFT+RANSAC", style='Caption')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph() # Spacer

        doc.add_page_break()

        # --- Visualisasi Pendukung (Plots) ---
        doc.add_heading('Visualisasi & Analisis Pendukung', level=1)
        
        plot_map = {
            'enhanced_localization_map': "Peta Detail Lokalisasi Tampering",
            'anomaly_infographic': "Infografis Penjelasan Anomali",
            'kmeans_temporal': "Visualisasi Temporal K-Means",
            'ssim_temporal': "Analisis Temporal SSIM",
            'optical_flow_temporal': "Analisis Temporal Aliran Optik",
            'temporal': "Peta Anomali Temporal Ringkas",
        }

        for key, caption_text in plot_map.items():
            plot_path_str = result.plots.get(key)
            if plot_path_str and Path(plot_path_str).exists():
                doc.add_heading(caption_text, level=2)
                try:
                    doc.add_picture(plot_path_str, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = doc.add_paragraph(f"Grafik: {caption_text}.", style='Caption')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception as img_err:
                    print(f"Warning: Gagal menambahkan gambar {plot_path_str} ke DOCX: {img_err}")

        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"FATAL: Terjadi error saat membuat file DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None


def create_docx_backend(result: Any, output_path: Path) -> Optional[Path]:
    """Wrapper kompatibilitas untuk pembuatan laporan DOCX."""
    return create_docx_report_robust(result, output_path)


def add_export_buttons(pdf_path: Path, result: Any, col1, col2, col3):
    """
    Menampilkan tombol ekspor dengan UI yang lebih sederhana dan fokus.
    Kolom PNG dihilangkan untuk saat ini.
    """
    # Kolom 1: Tombol Download PDF (Wajib)
    with col1:
        if pdf_path.exists():
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label="📄 Download PDF",
                    data=f.read(),
                    file_name=pdf_path.name,
                    mime="application/pdf",
                    use_container_width=True
                )
        else:
            st.error("Laporan PDF tidak berhasil dibuat. Periksa log konsol.")

    # Kolom 2: Tombol Download DOCX
    with col2:
        # Cari file docx di dalam atribut result
        docx_path_str = getattr(result, 'docx_report_path', None)
        if docx_path_str and Path(docx_path_str).exists():
             with open(docx_path_str, "rb") as f:
                st.download_button(
                    label="📝 Download DOCX",
                    data=f.read(),
                    file_name=Path(docx_path_str).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.warning("DOCX tidak dibuat. Cek log konsol untuk detailnya.", icon="⚠️")

    # Kolom 3 Dibiarkan kosong atau untuk penggunaan lain.
    with col3:
        st.write("") # Placeholder